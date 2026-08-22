import asyncio
from contextlib import contextmanager
from io import BytesIO
from pathlib import Path
from types import SimpleNamespace
from docx import Document
from fastapi import UploadFile
from fastapi.testclient import TestClient
from openpyxl import Workbook
from pypdf import PdfWriter
import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool
from app.core.database import Base,get_db
from app.main import app
from app.models import BackgroundJob,KnowledgeDocument,ModelCallLog,ModelProfile,Project,PromptTemplateVersion,RetrievalLog
from app.services.db.dialect import qualify_table,quote_identifier
from app.services.embeddings.factory import get_embedding_service
from app.services.embeddings.observability import embed_with_observability
from app.services.embeddings.openai_compatible import LocalEmbeddingService,OpenAICompatibleEmbeddingService
from app.services.rag.citation_validator import validate_citations
from app.services.llm.base import ModelCallMetadata
from app.services.llm.prompt_runtime import get_prompt_runtime,prepare_model_input,record_model_call
from app.services.knowledge_ingestion import ingest_knowledge_document
from app.services.security import ensure_external_allowed, redact_content
from app.services.vector import VectorRecord
from app.services.vector.factory import get_vector_store
from app.services.vector.milvus import MilvusVectorStore
from app.services.knowledge_ingestion.parsers import parse_document
from app.services.vector.mock import MockVectorStore
from app.services.storage.factory import get_storage_service
from app.services.task_queue.factory import get_task_queue
from app.services.retrieval.hybrid_retriever import _keyword_score
from app.services.retrieval.keyword_index import tokenize


def test_plain_text_parser_splits_windows_crlf_paragraphs() -> None:
    units, warnings = parse_document(
        "synthetic.md",
        "第一段\r\n\r\n第二段\r\n\r\n第三段".encode("utf-8"),
        "manual_note",
    )

    assert warnings == []
    assert [unit.content for unit in units] == ["第一段", "第二段", "第三段"]


def test_embedding_observability_uses_query_encoder_for_retrieval(db_session) -> None:
    class QueryAwareEmbedding:
        local_only = True
        last_call = ModelCallMetadata(provider="local_vllm", model="fake-query")

        def embed_query(self, text):
            return [9.0, 0.0]

        def embed_texts(self, texts):
            raise AssertionError("retrieval query must not use document encoding")

    vectors = embed_with_observability(
        db_session,
        1,
        QueryAwareEmbedding(),
        ["贷款余额如何计算"],
        ["internal"],
        input_type="query",
    )

    assert vectors == [[9.0, 0.0]]


def test_keyword_score_prioritizes_exact_field_identifiers_over_generic_bigrams() -> None:
    tokens = tokenize("LOAN_DTL.BALANCE_AMT 对应的业务含义是什么？")
    exact = SimpleNamespace(
        title=None,
        normalized_content="贷款余额来自 LOAN_DTL.BALANCE_AMT",
        target_field_code=None,
        scenario_id=None,
    )
    generic = SimpleNamespace(
        title=None,
        normalized_content="数据日期的业务规则",
        target_field_code=None,
        scenario_id=None,
    )

    assert _keyword_score(exact, tokens, None, None) > _keyword_score(
        generic,
        tokens,
        None,
        None,
    )


def test_knowledge_ingestion_reports_progress_by_batch(db_session, tmp_path: Path, monkeypatch):
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path))
    get_storage_service.cache_clear()
    get_embedding_service.cache_clear()
    get_vector_store.cache_clear()
    project = Project(name="批量知识索引")
    db_session.add(project)
    db_session.commit()
    db_session.refresh(project)
    workbook = Workbook()
    sheet = workbook.active
    sheet.append(["问题", "监管回复"])
    for index in range(5):
        sheet.append([f"问题 {index}", f"答复 {index}"])
    stream = BytesIO()
    workbook.save(stream)
    progress = []

    document = asyncio.run(
        ingest_knowledge_document(
            db_session,
            project.id,
            UploadFile(file=BytesIO(stream.getvalue()), filename="批量答疑.xlsx"),
            "regulatory_qa",
            batch_size=2,
            progress=lambda completed, total: progress.append((completed, total)),
        )
    )

    assert document.parse_summary_json["unit_count"] == 5
    assert progress == [(2, 5), (4, 5), (5, 5)]


def test_failed_knowledge_ingestion_with_same_hash_can_be_retried(
    db_session,
    tmp_path: Path,
    monkeypatch,
):
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path))
    get_storage_service.cache_clear()
    get_embedding_service.cache_clear()
    project = Project(name="失败索引重试")
    db_session.add(project)
    db_session.commit()
    db_session.refresh(project)
    store = MockVectorStore()
    original_upsert = store.upsert
    attempts = 0

    def fail_once(records):
        nonlocal attempts
        attempts += 1
        if attempts == 1:
            raise RuntimeError("temporary vector failure")
        return original_upsert(records)

    monkeypatch.setattr(store, "upsert", fail_once)
    monkeypatch.setattr(
        "app.services.knowledge_ingestion.ingestion_service.get_vector_store",
        lambda: store,
    )
    content = "retryable knowledge".encode()
    with pytest.raises(RuntimeError, match="temporary vector failure"):
        asyncio.run(
            ingest_knowledge_document(
                db_session,
                project.id,
                UploadFile(file=BytesIO(content), filename="retry.txt"),
                "manual_note",
            )
        )

    failed = db_session.query(KnowledgeDocument).one()
    assert failed.document_status == "failed"
    recovered = asyncio.run(
        ingest_knowledge_document(
            db_session,
            project.id,
            UploadFile(file=BytesIO(content), filename="retry.txt"),
            "manual_note",
        )
    )

    assert recovered.document_status == "indexed"
    assert recovered.current_version_no == 2
    assert attempts == 2


def test_version_switch_keeps_old_vectors_when_database_commit_fails(
    db_session,
    tmp_path: Path,
    monkeypatch,
):
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path))
    get_storage_service.cache_clear()
    get_embedding_service.cache_clear()
    project = Project(name="知识换版事务")
    db_session.add(project)
    db_session.commit()
    db_session.refresh(project)
    store = MockVectorStore()
    monkeypatch.setattr(
        "app.services.knowledge_ingestion.ingestion_service.get_vector_store",
        lambda: store,
    )
    first = asyncio.run(
        ingest_knowledge_document(
            db_session,
            project.id,
            UploadFile(file=BytesIO(b"first version"), filename="versioned.txt"),
            "manual_note",
        )
    )
    old_vector_id = f"knowledge-unit-{next(iter(store._records.values())).metadata['knowledge_unit_id']}"
    original_commit = db_session.commit
    commit_calls = 0

    def fail_final_commit():
        nonlocal commit_calls
        commit_calls += 1
        if commit_calls == 3:
            raise RuntimeError("final database commit failed")
        return original_commit()

    monkeypatch.setattr(db_session, "commit", fail_final_commit)
    with pytest.raises(RuntimeError, match="final database commit failed"):
        asyncio.run(
            ingest_knowledge_document(
                db_session,
                project.id,
                UploadFile(file=BytesIO(b"second version"), filename="versioned.txt"),
                "manual_note",
            )
        )

    db_session.refresh(first)
    assert first.document_status == "failed"
    assert old_vector_id in store._records


def test_knowledge_upload_returns_accepted_background_job_for_async_queue(tmp_path: Path, monkeypatch):
    class AsyncQueue:
        def enqueue(self, db, **kwargs):
            job = BackgroundJob(
                institution_id=kwargs["institution_id"],
                project_id=kwargs["project_id"],
                idempotency_key="async-knowledge-upload",
                job_type=kwargs["job_type"],
                status="queued",
                progress=0,
                payload_summary_json=kwargs["payload_summary"],
                result_summary_json={},
                created_by=kwargs["created_by"],
            )
            db.add(job)
            db.commit()
            db.refresh(job)
            return job

    monkeypatch.setenv("STORAGE_DIR", str(tmp_path))
    monkeypatch.setattr("app.services.task_queue.submission.get_task_queue", lambda: AsyncQueue())
    get_storage_service.cache_clear()
    with _client() as client:
        project = _post(client, "/api/projects", {"name": "异步知识上传"})
        response = client.post(
            f"/api/projects/{project['id']}/knowledge/documents/upload",
            data={
                "knowledge_type": "manual_note",
                "knowledge_scope": "project",
                "confidentiality_level": "internal",
            },
            files={"file": ("异步知识.txt", b"asynchronous knowledge", "text/plain")},
        )

    assert response.status_code == 202
    assert response.json()["job_type"] == "knowledge_ingestion"
    assert response.json()["status"] == "queued"


def test_knowledge_versions_hybrid_search_grounded_answer_and_evaluation(tmp_path:Path,monkeypatch):
    monkeypatch.setenv("STORAGE_DIR",str(tmp_path));get_vector_store.cache_clear();get_embedding_service.cache_clear()
    with _client() as client:
        project=_post(client,"/api/projects",{"name":"知识项目","bank_name":"甲银行"})
        payload=_qa_excel("监管答疑：客户证件类型应来自 ECIF_CUSTOMER.CERT_TYPE")
        document=_upload(client,project["id"],"监管答疑.xlsx",payload,"regulatory_qa","institution","甲银行")
        repeated=_upload(client,project["id"],"监管答疑.xlsx",payload,"regulatory_qa","institution","甲银行")
        assert repeated["id"]==document["id"];assert len(_get(client,f"/api/knowledge/documents/{document['id']}/versions?project_id={project['id']}"))==1
        units=_get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={document['id']}");assert len(units)==1;unit=units[0];assert unit["unit_type"]=="qa";assert unit["source_sheet_name"]=="答疑";assert unit["source_cell_range"]=="A2:F2"
        search=_post(client,f"/api/projects/{project['id']}/knowledge/hybrid-search",{"query":"客户证件类型 CERT_TYPE","knowledge_types":["regulatory_qa"],"top_k":5});assert search["items"][0]["knowledge_unit_id"]==unit["id"];assert search["items"][0]["source_cell_range"]=="A2:F2"
        answer=_post(client,f"/api/projects/{project['id']}/knowledge/ask",{"query":"客户证件类型取哪个字段","top_k":5});assert answer["citations"][0]["knowledge_unit_id"]==unit["id"]
        case=_post(client,f"/api/projects/{project['id']}/evaluations/cases",{"case_name":"证件类型召回","query_text":"客户证件类型 CERT_TYPE","expected_knowledge_unit_ids_json":[unit["id"]],"expected_source_system":"ECIF","expected_table_name":"ECIF_CUSTOMER","expected_field_name":"CERT_TYPE","expected_answer_keywords_json":["CERT_TYPE"]});run=_post(client,f"/api/projects/{project['id']}/evaluations/runs",{"run_name":"回归"});metrics=run["summary_metrics_json"];assert metrics["recall_at_5"]==1;assert metrics["recall_at_10"]==1;assert metrics["mrr"]==1;assert metrics["source_hit_rate"]==1;assert metrics["table_hit_rate"]==1;assert metrics["field_hit_rate"]==1;assert metrics["citation_coverage"]==1;assert metrics["keyword_coverage"]==1
        feedback=_post(client,f"/api/projects/{project['id']}/feedback",{"feedback_type":"retrieval","target_type":"knowledge_unit","target_id":unit["id"],"rating":"correct","comment":"引用准确"});assert feedback["rating"]=="correct"
        changed=_upload(client,project["id"],"监管答疑.xlsx",_qa_excel("更新答疑：优先取 ECIF_CUSTOMER.CERT_TYPE"),"regulatory_qa","institution","甲银行");assert changed["id"]==document["id"];assert len(_get(client,f"/api/knowledge/documents/{document['id']}/versions?project_id={project['id']}"))==2;assert len(_get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={document['id']}"))==1
        other=_post(client,"/api/projects",{"name":"隔离项目","bank_name":"乙银行"});isolated=_post(client,f"/api/projects/{other['id']}/knowledge/hybrid-search",{"query":"CERT_TYPE","top_k":5});assert isolated["items"]==[]
        assert client.get(f"/api/knowledge/documents/{document['id']}?project_id={other['id']}").status_code==404
        assert client.get(f"/api/knowledge/units/{unit['id']}?project_id={other['id']}").status_code==404
        assert client.delete(f"/api/knowledge/documents/{document['id']}?project_id={other['id']}").status_code==404

def test_docx_pdf_text_markdown_and_sql_ingestion_preserve_locations(tmp_path:Path,monkeypatch):
    monkeypatch.setenv("STORAGE_DIR",str(tmp_path));get_vector_store.cache_clear();get_embedding_service.cache_clear()
    with _client() as client:
        project=_post(client,"/api/projects",{"name":"多格式知识"})
        document=Document();document.add_heading("第一章 监管范围",level=1);document.add_paragraph("本条规定客户证件类型的适用范围。");table=document.add_table(rows=1,cols=2);table.cell(0,0).text="字段";table.cell(0,1).text="CERT_TYPE";stream=BytesIO();document.save(stream)
        docx=_upload(client,project["id"],"制度.docx",stream.getvalue(),"regulatory_policy");docx_units=_get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={docx['id']}");assert any(item["source_heading"]=="第一章 监管范围" for item in docx_units);assert any(item["source_cell_range"]=="table:1:row:1" for item in docx_units)
        writer=PdfWriter();writer.add_blank_page(width=200,height=200);pdf=BytesIO();writer.write(pdf);pdf_doc=_upload(client,project["id"],"说明.pdf",pdf.getvalue(),"regulatory_policy");assert "第 1 页" in str(pdf_doc["warnings_json"])
        txt=_upload(client,project["id"],"调研.txt","业务调研结论：借记卡客户范围待确认。".encode(),"business_research");assert _get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={txt['id']}")[0]["unit_type"]=="paragraph"
        md=_upload(client,project["id"],"解释.md","## 字段解释\n\n客户证件类型代码。".encode(),"field_explanation");assert _get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={md['id']}")
        sql=_upload(client,project["id"],"历史.sql",b"select cert_type, case when cert_type='01' then 'ID' end from ecif_customer where status='A'","sql_evidence");sql_unit=_get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={sql['id']}")[0];assert "ecif_customer" in sql_unit["content"];assert "cert_type" in sql_unit["content"]

def test_dialect_identifiers_and_datasource_guard():
    assert quote_identifier("order","mysql")=="`order`";assert qualify_table("ODS","CUSTOMER","mysql_compatible")=="`ODS`.`CUSTOMER`";assert qualify_table("main","customer","sqlite")=='"customer"'


def test_scope_reuse_isolation_and_no_evidence_answer(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path))
    get_vector_store.cache_clear()
    get_embedding_service.cache_clear()
    with _client() as client:
        owner = _post(client, "/api/projects", {"name": "甲行项目一", "bank_name": "甲银行"})
        same_bank = _post(client, "/api/projects", {"name": "甲行项目二", "bank_name": "甲银行"})
        other_bank = _post(client, "/api/projects", {"name": "乙行项目", "bank_name": "乙银行"})
        _upload(client, owner["id"], "通用制度.txt", "监管通用证件类型规则".encode(), "regulatory_policy", "global")
        _upload(client, owner["id"], "甲行口径.txt", "甲银行内部证件类型规则".encode(), "historical_mapping", "institution", "甲银行")

        same_result = _post(client, f"/api/projects/{same_bank['id']}/knowledge/hybrid-search", {"query": "证件类型规则", "top_k": 10})
        # bank_name is descriptive text, not an institution authorization key.
        assert {item["knowledge_type"] for item in same_result["items"]} == {"regulatory_policy"}
        assert all(item["vector_score"] > 0 for item in same_result["items"])

        other_result = _post(client, f"/api/projects/{other_bank['id']}/knowledge/hybrid-search", {"query": "证件类型规则", "top_k": 10})
        assert {item["knowledge_type"] for item in other_result["items"]} == {"regulatory_policy"}
        no_evidence = _post(client, f"/api/projects/{other_bank['id']}/knowledge/ask", {"query": "完全不存在的量子字段 ZXQ_999", "knowledge_types": ["sql_evidence"], "top_k": 5})
        assert no_evidence["citations"] == []
        assert "待确认" in no_evidence["answer"]


def test_redaction_external_policy_and_citation_validation(db_session):
    raw = "手机号 13800138000，证件号 110101199001011234，邮箱 user@example.com，password=secret"
    redacted = redact_content(raw)
    assert "13800138000" not in redacted
    assert "110101199001011234" not in redacted
    assert "user@example.com" not in redacted
    assert "secret" not in redacted
    with pytest.raises(ValueError, match="restricted"):
        ensure_external_allowed("restricted", local_only=False)
    with pytest.raises(ValueError, match="confidential"):
        ensure_external_allowed("confidential", local_only=False)
    validate_citations(db_session, [])
    with pytest.raises(ValueError, match="不存在"):
        validate_citations(db_session, [{"knowledge_unit_id": 999999}])


def test_openai_compatible_embedding_and_milvus_adapter(monkeypatch):
    class Response:
        def __init__(self, count):
            self.count = count

        def raise_for_status(self):
            return None

        def json(self):
            values = [[0.1, 0.2], [0.3, 0.4]]
            return {"data": [{"embedding": value} for value in values[: self.count]]}

    calls = []

    def fake_post(url, **kwargs):
        calls.append((url, kwargs))
        return Response(len(kwargs["json"]["input"]))

    monkeypatch.setenv("TEST_EMBEDDING_KEY", "not-a-real-secret")
    monkeypatch.setattr("app.services.embeddings.openai_compatible.httpx.post", fake_post)
    embedding = OpenAICompatibleEmbeddingService("http://embedding.test/v1", "demo-model", "TEST_EMBEDDING_KEY")
    assert embedding.embed_texts(["a", "b"]) == [[0.1, 0.2], [0.3, 0.4]]
    assert calls[0][0] == "http://embedding.test/v1/embeddings"
    assert calls[0][1]["json"] == {"model": "demo-model", "input": ["a", "b"]}

    calls.clear()
    local_embedding = LocalEmbeddingService(
        "http://127.0.0.1:11434/v1",
        "BAAI/bge-small-zh-v1.5",
        "",
    )
    assert local_embedding.embed_query("贷款余额") == [0.1, 0.2]
    assert calls[0][1]["headers"]["X-YBT-Embedding-Input-Type"] == "query"
    assert "input_type" not in calls[0][1]["json"]

    class FakeMilvusClient:
        def __init__(self):
            self.created = []
            self.upserts = []
            self.searches = []
            self.deletes = []

        def has_collection(self, name):
            return False

        def create_collection(self, **kwargs):
            self.created.append(kwargs)

        def upsert(self, collection_name, data):
            self.upserts.append((collection_name, data))

        def search(self, collection_name, vectors, **kwargs):
            self.searches.append((collection_name, vectors, kwargs))
            return [[{"id": "knowledge-unit-1", "distance": 0.9, "entity": {"content": "证据", "project_id": 1}}]]

        def delete(self, collection_name, **kwargs):
            self.deletes.append((collection_name, kwargs))

    client = FakeMilvusClient()
    store = MilvusVectorStore(client=client)
    store.upsert([VectorRecord("knowledge-unit-1", [0.1, 0.2], "证据", {"project_id": 1})])
    result = store.search([0.1, 0.2], 5, {"project_id": 1, "knowledge_type": ["regulatory_qa"]})
    store.delete(filters={"project_id": 1})
    assert result[0].metadata["project_id"] == 1
    assert client.created[0]["dimension"] == 2
    assert 'knowledge_type in ["regulatory_qa"]' in client.searches[0][2]["filter"]
    assert client.deletes[0][1]["filter"] == "project_id == 1"


def test_ingestion_and_reindex_never_store_raw_sensitive_content_in_milvus(tmp_path:Path,monkeypatch):
    class RecordingMilvusClient:
        def __init__(self):self.upserts=[]
        def has_collection(self,name):return True
        def upsert(self,collection_name,data):self.upserts.extend(data)
        def delete(self,collection_name,**kwargs):return None
    client=RecordingMilvusClient();store=MilvusVectorStore(client=client)
    monkeypatch.setenv("STORAGE_DIR",str(tmp_path));get_embedding_service.cache_clear();get_vector_store.cache_clear()
    monkeypatch.setattr("app.services.knowledge_ingestion.ingestion_service.get_vector_store",lambda:store)
    monkeypatch.setattr("app.api.knowledge_rag.get_vector_store",lambda:store)
    raw="restricted-original-phrase 客户手机号 13800138000 证件号 110101199001011234 账号 6222020202020202 password=plain-secret"
    with _client() as api:
        project=_post(api,"/api/projects",{"name":"向量正文安全"})
        document=_upload(api,project["id"],"受限知识.txt",raw.encode(),"manual_note",confidentiality="restricted")
        _assert_milvus_payload_redacted(client.upserts,raw)
        client.upserts.clear()
        _post(api,f"/api/knowledge/documents/{document['id']}/reindex?project_id={project['id']}",{})
        _assert_milvus_payload_redacted(client.upserts,raw)


def test_milvus_count_prefers_live_aggregate_over_lagging_collection_stats():
    class LaggingStatsClient:
        def has_collection(self, name):
            return True

        def query(self, collection_name, filter, output_fields):
            assert filter == ""
            assert output_fields == ["count(*)"]
            return [{"count(*)": 100}]

        def get_collection_stats(self, collection_name):
            return {"row_count": 0}

    store = MilvusVectorStore(client=LaggingStatsClient(), collection_name="semantic_v1")

    assert store.count() == 100


def test_milvus_validation_flushes_pending_writes_before_counting():
    class FlushAwareClient:
        flushed = False

        def has_collection(self, name):
            return True

        def flush(self, collection_name):
            self.flushed = True

        def query(self, collection_name, filter, output_fields):
            return [{"count(*)": 100 if self.flushed else 0}]

        def get_collection_stats(self, collection_name):
            return {"row_count": 0}

    client = FlushAwareClient()
    store = MilvusVectorStore(
        client=client,
        collection_name="semantic_v1",
        expected_dimension=512,
    )

    result = store.validate_index(expected_count=100, expected_dimension=512)

    assert client.flushed is True
    assert result["valid"] is True
    assert result["actual_count"] == 100


def test_hybrid_retriever_revalidates_database_scope_and_enabled_for_stale_vectors(tmp_path:Path,monkeypatch):
    class StaleVectorStore(MockVectorStore):
        def delete(self,ids=None,filters=None):return None
    store=StaleVectorStore();monkeypatch.setenv("STORAGE_DIR",str(tmp_path));get_embedding_service.cache_clear();get_vector_store.cache_clear()
    monkeypatch.setattr("app.services.knowledge_ingestion.ingestion_service.get_vector_store",lambda:store)
    monkeypatch.setattr("app.api.knowledge_rag.get_vector_store",lambda:store)
    monkeypatch.setattr("app.services.retrieval.hybrid_retriever.get_vector_store",lambda:store)
    with _client() as api:
        owner=_post(api,"/api/projects",{"name":"甲行向量项目","bank_name":"甲银行"});other=_post(api,"/api/projects",{"name":"乙行向量项目","bank_name":"乙银行"})
        document=_upload(api,owner["id"],"银行知识.txt",b"STALE_VECTOR_SECRET","manual_note","institution","甲银行")
        unit=_get(api,f"/api/projects/{owner['id']}/knowledge/units?document_id={document['id']}")[0]
        vector=get_embedding_service().embed_query("STALE_VECTOR_SECRET")
        metadata={"project_id":owner["id"],"knowledge_scope":"institution","institution_name":"乙银行","knowledge_type":"manual_note","knowledge_unit_id":unit["id"]}
        store.upsert([VectorRecord(f"knowledge-unit-{unit['id']}",vector,"",metadata)])
        cross_bank=_post(api,f"/api/projects/{other['id']}/knowledge/hybrid-search",{"query":"STALE_VECTOR_SECRET","top_k":5})
        assert cross_bank["items"]==[]
        metadata["institution_name"]="甲银行";store.upsert([VectorRecord(f"knowledge-unit-{unit['id']}",vector,"",metadata)])
        response=api.delete(f"/api/knowledge/documents/{document['id']}?project_id={owner['id']}");assert response.status_code==200,response.text
        disabled=_post(api,f"/api/projects/{owner['id']}/knowledge/hybrid-search",{"query":"STALE_VECTOR_SECRET","top_k":5})
        assert disabled["items"]==[]


def test_model_profile_rejects_nested_credentials(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path))
    with _client() as client:
        response = client.post("/api/model-profiles", json={
            "profile_name": "unsafe-profile",
            "provider_type": "openai_compatible",
            "config_json": {"headers": {"AuthorizationToken": "plaintext-secret"}},
        })
        assert response.status_code == 400
        assert "credentials" in response.text


def test_excel_merged_cells_scenario_links_and_persistent_keyword_index(tmp_path:Path,monkeypatch):
    monkeypatch.setenv("STORAGE_DIR",str(tmp_path));get_vector_store.cache_clear();get_embedding_service.cache_clear()
    with _client() as client:
        project=_post(client,"/api/projects",{"name":"场景知识索引"})
        table=_post(client,"/api/target-tables",{"project_id":project["id"],"table_code":"YBT_CUSTOMER","table_name":"客户"})
        _post(client,"/api/fields",{"project_id":project["id"],"target_table_id":table["id"],"field_code":"CERT_TYPE","field_name":"客户证件类型"})
        debit=_post(client,f"/api/projects/{project['id']}/scenarios",{"scenario_code":"DEBIT","scenario_name":"借记卡"})
        credit=_post(client,f"/api/projects/{project['id']}/scenarios",{"scenario_code":"CREDIT","scenario_name":"信用卡"})
        workbook=Workbook();sheet=workbook.active;sheet.title="历史场景";sheet.append(["业务场景","字段代码","字段名称","来源表","来源字段","备注"]);sheet.append(["借记卡","CERT_TYPE","客户证件类型","ecif_customer","cert_type","合并备注"]);sheet.append(["信用卡","CERT_TYPE","客户证件类型","ecif_customer","cert_type",None]);sheet.merge_cells("F2:F3");stream=BytesIO();workbook.save(stream)
        document=_upload(client,project["id"],"历史场景.xlsx",stream.getvalue(),"historical_mapping")
        units=_get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={document['id']}")
        assert {item["scenario_id"] for item in units}=={debit["id"],credit["id"]}
        assert all("合并备注" in item["content"] for item in units)
        paragraphs=[f"普通历史知识段落 {index}" for index in range(550)]+["持久关键词索引唯一标记 UNIQUE_LAST_550"]
        _upload(client,project["id"],"大量知识.txt","\n\n".join(paragraphs).encode(),"manual_note")
        result=_post(client,f"/api/projects/{project['id']}/knowledge/hybrid-search",{"query":"UNIQUE_LAST_550","top_k":5})
        assert result["items"] and "UNIQUE_LAST_550" in result["items"][0]["content"]


def test_prompt_version_external_policy_and_model_call_audit(db_session):
    project=Project(name="Prompt 审计项目");db_session.add(project);db_session.flush()
    profile=ModelProfile(profile_name="external",provider_type="openai_compatible",base_url="https://model.example/v1",model_name="demo",enabled=True,local_only=False,config_json={"api_key_env_name":"TEST_MODEL_KEY"});db_session.add(profile)
    prompt=PromptTemplateVersion(prompt_key="scenario_business_mapping",version_no=7,system_prompt="场景业务口径",user_prompt_template="{evidence}",enabled=True);db_session.add(prompt)
    retrieval=RetrievalLog(project_id=project.id,query_text="证件类型",query_type="hybrid",filters_json={},retrieval_strategy="test",result_ids_json=[]);db_session.add(retrieval);db_session.commit()
    runtime=get_prompt_runtime(db_session,"scenario_business_mapping")
    assert runtime.version==7 and runtime.model_profile_id==profile.id and runtime.local_only is False
    with pytest.raises(ValueError,match="restricted"):
        prepare_model_input(runtime,"受限知识",["restricted"])
    model_input=prepare_model_input(runtime,"联系电话 13800138000",["internal"])
    assert "13800138000" not in model_input
    record_model_call(db_session,project.id,runtime,model_input,{"draft":"联系电话 13800138000"},retrieval_log_id=retrieval.id)
    db_session.commit();log=db_session.query(ModelCallLog).one()
    assert log.prompt_version==7 and log.retrieval_log_id==retrieval.id
    assert "13800138000" not in (log.output_summary or "")

def _qa_excel(answer):
    workbook=Workbook();sheet=workbook.active;sheet.title="答疑";sheet.append(["问题","监管回复","表代码","字段代码","字段名称","备注"]);sheet.append(["客户证件类型如何取值",answer,"YBT_CUSTOMER","CERT_TYPE","客户证件类型","脱敏模拟"]);stream=BytesIO();workbook.save(stream);return stream.getvalue()
def _upload(client,project_id,name,content,kind,scope="project",institution=None,confidentiality="internal"):
    response=client.post(f"/api/projects/{project_id}/knowledge/documents/upload",data={"knowledge_type":kind,"knowledge_scope":scope,"institution_name":institution or "","confidentiality_level":confidentiality},files={"file":(name,content,"application/octet-stream")});assert response.status_code==200,response.text;return response.json()

def _assert_milvus_payload_redacted(rows,raw):
    payload=str(rows)
    for forbidden in [raw,"restricted-original-phrase","13800138000","110101199001011234","6222020202020202","plain-secret"]:assert forbidden not in payload
    allowed={"id","vector","knowledge_unit_id","project_id","knowledge_scope","institution_name","knowledge_type","target_field_code","scenario_id","confidentiality_level","document_version_id","content_hash"}
    assert rows and all(set(row)<=allowed and "content" not in row for row in rows)
def _post(client,path,payload):
    response=client.post(path,json=payload);assert response.status_code==200,response.text;return response.json()
def _get(client,path):
    response=client.get(path);assert response.status_code==200,response.text;return response.json()
@contextmanager
def _client():
    get_storage_service.cache_clear();get_task_queue.cache_clear()
    engine=create_engine("sqlite://",connect_args={"check_same_thread":False},poolclass=StaticPool);Base.metadata.create_all(engine);factory=sessionmaker(bind=engine,autoflush=False)
    def override():
        session=factory()
        try:yield session
        finally:session.close()
    app.dependency_overrides[get_db]=override
    try:
        with TestClient(app) as client:yield client
    finally:app.dependency_overrides.clear();Base.metadata.drop_all(engine);get_vector_store.cache_clear();get_embedding_service.cache_clear();get_storage_service.cache_clear();get_task_queue.cache_clear()
