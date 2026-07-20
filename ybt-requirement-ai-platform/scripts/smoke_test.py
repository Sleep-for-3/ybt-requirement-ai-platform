import json
import sqlite3
import tempfile
from io import BytesIO
from pathlib import Path

import httpx
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font


def main() -> None:
    base = "http://127.0.0.1:8000/api"
    client = httpx.Client(timeout=90, trust_env=False)

    admin_password = "smoke-only-" + "platform-admin-password"
    bootstrap = _post_json(client, f"{base}/admin/bootstrap", {
        "institution_code": "PLATFORM", "institution_name": "脱敏平台运营方", "institution_type": "platform_operator",
        "username": "smoke_admin", "display_name": "Smoke 管理员", "email": "smoke-admin@example.invalid", "password": admin_password,
    })
    admin_session = _post_json(client, f"{base}/auth/login", {"username": "smoke_admin", "password": admin_password})
    client.headers["Authorization"] = f"Bearer {admin_session['access_token']}"
    bank = _post_json(client, f"{base}/admin/institutions", {"institution_code": "DEMO_BANK", "institution_name": "示例银行", "institution_type": "bank"})

    project = _post_json(
        client,
        f"{base}/projects",
        {
            "name": "增强验收测试项目",
            "institution_id": bank["id"],
            "bank_name": "示例银行",
            "description": "验证模板、数据源、自然语言任务和口径生成",
        },
    )
    project_id = project["id"]

    with tempfile.TemporaryDirectory() as temp_dir:
        temp_path = Path(temp_dir)
        excel_path = temp_path / "一表通模板.xlsx"
        traceability_path = temp_path / "脱敏业务口径及溯源表.xlsx"
        qa_path = temp_path / "监管答疑.xlsx"
        historical_path = temp_path / "历史业务口径.xlsx"
        policy_path = temp_path / "监管制度.docx"
        pdf_path = temp_path / "监管说明.pdf"
        sql_path = temp_path / "历史取数逻辑.sql"
        sqlite_path = temp_path / "ecif_query.db"
        _write_template(excel_path)
        _write_traceability_template(traceability_path)
        _write_qa(qa_path)
        _write_historical_mapping(historical_path)
        _write_policy_docx(policy_path)
        _write_text_pdf(pdf_path)
        sql_path.write_text("select cert_type from ecif_customer where status = 'A'", encoding="utf-8")
        _write_sqlite_source(sqlite_path)

        with excel_path.open("rb") as file:
            template = _post_file(
                client,
                f"{base}/templates/upload",
                data={"project_id": str(project_id)},
                files={"file": ("一表通模板.xlsx", file, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            )
        apply_result = _post_json(client, f"{base}/templates/{template['template_id']}/apply", {})

        with traceability_path.open("rb") as file:
            traceability_template = _post_file(
                client,
                f"{base}/traceability-templates/upload",
                data={"project_id": str(project_id)},
                files={"file": ("脱敏业务口径及溯源表.xlsx", file, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            )
        traceability_preview = _get_json(client, f"{base}/traceability-templates/{traceability_template['template_id']}/preview")
        traceability_apply = _post_json(client, f"{base}/traceability-templates/{traceability_template['template_id']}/apply", {})

        fields = client.get(f"{base}/fields", params={"project_id": project_id})
        fields.raise_for_status()
        field = next(item for item in fields.json() if item["field_code"] == "CERT_TYPE")

        knowledge_documents = [
            _upload_knowledge(client, base, project_id, qa_path, "regulatory_qa", "institution", "示例银行"),
            _upload_knowledge(client, base, project_id, historical_path, "historical_mapping"),
            _upload_knowledge(client, base, project_id, policy_path, "regulatory_policy"),
            _upload_knowledge(client, base, project_id, pdf_path, "regulatory_policy"),
            _upload_knowledge(client, base, project_id, sql_path, "sql_evidence"),
        ]
        knowledge_units = _get_json(client, f"{base}/projects/{project_id}/knowledge/units")
        qa_unit = next(item for item in knowledge_units if item["knowledge_type"] == "regulatory_qa")
        if qa_unit.get("source_sheet_name") != "答疑" or not qa_unit.get("source_cell_range"):
            raise AssertionError("监管答疑未保留 sheet 和单元格引用")
        pdf_unit = next(item for item in knowledge_units if item["source_file_name"] == "监管说明.pdf")
        if pdf_unit.get("source_page_no") != 1:
            raise AssertionError("PDF 知识未保留页码")
        knowledge_search = _post_json(client, f"{base}/projects/{project_id}/knowledge/hybrid-search", {
            "query": "借记卡客户证件类型 CERT_TYPE", "target_field_id": field["id"], "top_k": 10,
        })
        search_types = {item["knowledge_type"] for item in knowledge_search["items"]}
        if not {"regulatory_qa", "historical_mapping"}.issubset(search_types):
            raise AssertionError(f"混合检索未命中监管答疑和历史口径: {search_types}")
        grounded = _post_json(client, f"{base}/projects/{project_id}/knowledge/ask", {
            "query": "借记卡客户证件类型取哪个字段", "target_field_id": field["id"], "top_k": 10,
        })
        if not grounded["citations"] or grounded["citations"][0]["knowledge_unit_id"] not in {item["id"] for item in knowledge_units}:
            raise AssertionError("知识问答未返回真实 citation")
        no_evidence = _post_json(client, f"{base}/projects/{project_id}/knowledge/ask", {
            "query": "完全未知字段 ZXQ_999", "knowledge_types": ["manual_note"], "top_k": 5,
        })
        if no_evidence["citations"] or "待确认" not in no_evidence["answer"]:
            raise AssertionError("无证据问答未返回待确认")
        evaluation_case = _post_json(client, f"{base}/projects/{project_id}/evaluations/cases", {
            "case_name": "借记卡证件类型召回", "query_text": "借记卡客户证件类型 CERT_TYPE",
            "target_field_id": field["id"], "expected_knowledge_unit_ids_json": [qa_unit["id"]],
            "expected_table_name": "ecif_customer", "expected_field_name": "cert_type",
            "expected_answer_keywords_json": ["CERT_TYPE"],
        })
        evaluation_run = _post_json(client, f"{base}/projects/{project_id}/evaluations/runs", {"run_name": "Smoke RAG 回归"})
        if evaluation_run["summary_metrics_json"].get("recall_at_5", 0) <= 0 or evaluation_run["summary_metrics_json"].get("mrr", 0) <= 0:
            raise AssertionError("RAG 评测 Recall@5 或 MRR 未命中")

        datasource = _post_json(
            client,
            f"{base}/projects/{project_id}/datasources",
            {
                "name": "ecif_query",
                "display_name": "ECIF SQLite 测试库",
                "db_type": "sqlite",
                "database_name": str(sqlite_path),
                "readonly_flag": True,
                "enabled": True,
            },
        )
        datasource_test = _post_json(client, f"{base}/datasources/{datasource['id']}/test", {})
        metadata_sync = _post_json(client, f"{base}/datasources/{datasource['id']}/metadata-sync", {"sync_mode": "full", "schema_names": [], "include_views": True})
        catalog_tables = _get_json(client, f"{base}/projects/{project_id}/catalog/tables?datasource_id={datasource['id']}")
        catalog_search = _post_json(client, f"{base}/projects/{project_id}/catalog/search", {"datasource_ids": [datasource["id"]], "query": "cert_type", "top_k": 20})
        catalog_cert_type = next(item for item in catalog_search["items"] if item["column_name"] == "cert_type")
        catalog_source_import = _post_json(client, f"{base}/catalog/columns/{catalog_cert_type['catalog_column_id']}/import-as-source-field", {"system_code": "CATALOG_ECIF", "system_name": "目录导入客户系统"})

        task_create = _post_json(
            client,
            f"{base}/nl-tasks",
            {
                "project_id": project_id,
                "text": "使用 ecif_query 查询 ecif_customer 表 cert_type 字段的空值率和枚举分布",
            },
        )
        task_run = _post_json(client, f"{base}/nl-tasks/{task_create['task_id']}/run", {})

        business_system = _post_json(
            client,
            f"{base}/projects/{project_id}/business-systems",
            {"system_code": "ECIF", "system_name": "客户信息系统", "owner_department": "数据管理部", "enabled": True},
        )
        source_table = _post_json(
            client,
            f"{base}/business-systems/{business_system['id']}/source-tables",
            {
                "table_code": "ecif_customer",
                "table_name": "客户基本信息表",
                "table_comment": "ECIF 客户主表",
                "datasource_id": datasource["id"],
                "physical_table_name": "ecif_customer",
            },
        )
        source_field = _post_json(
            client,
            f"{base}/source-tables/{source_table['id']}/source-fields",
            {
                "field_code": "cert_type",
                "field_name": "证件类型",
                "field_type": "text",
                "field_comment": "客户证件类型",
                "physical_column_name": "cert_type",
            },
        )
        scenarios = _get_json(client, f"{base}/projects/{project_id}/scenarios")
        debit_scenario = next(item for item in scenarios if item["scenario_code"] == "DEBIT_CARD")
        business_mappings = _get_json(client, f"{base}/target-fields/{field['id']}/scenario-business-mappings")
        scenario_business = next(item for item in business_mappings if item["scenario_id"] == debit_scenario["id"])
        existing_lineages = _get_json(client, f"{base}/target-fields/{field['id']}/scenario-technical-lineages")
        existing_scenario_lineage = next(item for item in existing_lineages if item["scenario_id"] == debit_scenario["id"])
        source_recommendations = _post_json(
            client,
            f"{base}/target-fields/{field['id']}/scenarios/{debit_scenario['id']}/recommend-sources",
            {},
        )
        catalog_recommendation = next(item for item in source_recommendations["recommendations"] if item.get("catalog_column_id") == catalog_cert_type["catalog_column_id"])
        if not catalog_recommendation.get("knowledge_unit_ids_json") or not catalog_recommendation.get("citation_summary_json"):
            raise AssertionError("目录来源推荐未携带知识引用")
        selected_recommendation = _post_json(
            client,
            f"{base}/source-recommendations/{catalog_recommendation['id']}/select",
            {},
        )
        scenario_lineage = selected_recommendation["lineage"]
        if scenario_lineage.get("source_field_english_name") != existing_scenario_lineage.get("source_field_english_name"):
            raise AssertionError("目录候选仅选择时不应自动写入技术来源")
        column_profile = _post_json(client, f"{base}/catalog/columns/{catalog_cert_type['catalog_column_id']}/profile", {
            "target_field_id": field["id"], "scenario_id": debit_scenario["id"], "source_recommendation_id": catalog_recommendation["id"],
            "metrics": ["null_rate", "distinct_count", "top_values", "min_max", "length_distribution"],
        })
        if column_profile["profile_result_json"].get("distinct_count") != 2:
            raise AssertionError("目录字段 distinct 探查结果不正确")
        selected_recommendation = _post_json(client, f"{base}/source-recommendations/{catalog_recommendation['id']}/adopt", {})
        scenario_lineage = selected_recommendation["lineage"]

        sensitive_field = _post_json(client, f"{base}/fields", {"project_id": project_id, "target_table_id": field["target_table_id"], "field_code": "CUSTOMER_NAME", "field_name": "客户姓名"})
        sensitive_recommendations = _post_json(client, f"{base}/target-fields/{sensitive_field['id']}/scenarios/{debit_scenario['id']}/recommend-sources", {})
        sensitive_catalog = next(item for item in sensitive_recommendations["recommendations"] if item.get("catalog_column_id") and item.get("recommended_field_name") == "customer_name")
        _post_json(client, f"{base}/source-recommendations/{sensitive_catalog['id']}/select", {})
        sensitive_profile = _post_json(client, f"{base}/catalog/columns/{sensitive_catalog['catalog_column_id']}/profile", {"target_field_id": sensitive_field["id"], "scenario_id": debit_scenario["id"], "source_recommendation_id": sensitive_catalog["id"], "metrics": ["distinct_count", "top_values", "min_max"]})
        if sensitive_profile["profile_result_json"].get("top_values"):
            raise AssertionError("敏感字段不应返回 top values")

        manual_business = "人工确认：借记卡场景证件类型按有效借记卡客户业务定义维护。"
        _put_json(client, f"{base}/scenario-business-mappings/{scenario_business['id']}", {"final_content": manual_business})
        generated_business = _post_json(client, f"{base}/scenario-business-mappings/{scenario_business['id']}/generate-draft", {})
        if generated_business["final_content"] != manual_business:
            raise AssertionError("AI 场景业务草稿覆盖了人工 final_content")
        adopted_business = _post_json(client, f"{base}/scenario-business-mappings/{scenario_business['id']}/adopt-ai-draft", {})
        _post_json(client, f"{base}/mappings/scenario_business/{scenario_business['id']}/evidence", {
            "evidence_type": "manual_note", "source_name": "Smoke 脱敏业务访谈记录",
            "evidence_summary": "业务部门确认场景业务口径。",
        })
        legacy_business_confirm = client.post(f"{base}/scenario-business-mappings/{scenario_business['id']}/confirm", json={"confirmed_by": "smoke"})
        if legacy_business_confirm.status_code != 409:
            raise AssertionError(f"治理模式旧业务确认接口必须返回 409，实际为 {legacy_business_confirm.status_code}")

        manual_technical = "人工确认：来源为 ECIF.ecif_customer.cert_type，按借记卡有效客户范围取值。"
        _put_json(client, f"{base}/scenario-technical-lineages/{scenario_lineage['id']}", {"final_content": manual_technical})
        generated_lineage = _post_json(client, f"{base}/scenario-technical-lineages/{scenario_lineage['id']}/generate-draft", {})
        if generated_lineage["final_content"] != manual_technical:
            raise AssertionError("AI 场景技术草稿覆盖了人工 final_content")
        if "安全探查摘要" not in generated_lineage["ai_generated_content"]:
            raise AssertionError("AI 场景技术草稿未引用目录探查摘要")
        adopted_lineage = _post_json(client, f"{base}/scenario-technical-lineages/{scenario_lineage['id']}/adopt-ai-draft", {})
        _post_json(client, f"{base}/mappings/scenario_technical/{scenario_lineage['id']}/evidence", {
            "evidence_type": "manual_note", "source_name": "Smoke 脱敏技术访谈记录",
            "evidence_summary": "科技部门确认场景技术溯源。",
        })
        legacy_technical_confirm = client.post(f"{base}/scenario-technical-lineages/{scenario_lineage['id']}/confirm", json={"confirmed_by": "smoke"})
        if legacy_technical_confirm.status_code != 409:
            raise AssertionError(f"治理模式旧技术确认接口必须返回 409，实际为 {legacy_technical_confirm.status_code}")
        mart_table = _post_json(
            client,
            f"{base}/projects/{project_id}/mart-tables",
            {
                "table_code": "mart_customer",
                "table_name": "监管客户集市表",
                "subject_area": "客户",
                "table_comment": "一表通客户主题中间层",
                "is_existing": False,
            },
        )
        mart_field = _post_json(
            client,
            f"{base}/mart-tables/{mart_table['id']}/mart-fields",
            {
                "field_code": "cert_type",
                "field_name": "客户证件类型",
                "field_type": "varchar(20)",
                "field_comment": "统一监管客户证件类型",
                "is_existing": False,
            },
        )
        source_to_mart = _post_json(
            client,
            f"{base}/mart-fields/{mart_field['id']}/source-to-mart-mappings",
            {
                "mapping_name": "ECIF 证件类型入监管集市",
                "source_system_summary": "ECIF",
                "source_tables_summary": "ecif_customer",
                "source_fields_summary": "cert_type",
                "business_rule": "从 ECIF 客户基本信息表取客户证件类型。",
            },
        )
        mart_to_ybt = _post_json(
            client,
            f"{base}/target-fields/{field['id']}/mart-to-ybt-mappings",
            {
                "mart_field_id": mart_field["id"],
                "mapping_name": "监管集市证件类型到一表通",
                "mart_table_summary": "mart_customer",
                "mart_field_summary": "cert_type",
                "business_rule": "从监管客户集市字段取值并转换为一表通代码。",
            },
        )
        source_evidence = _post_json(
            client,
            f"{base}/mappings/source_to_mart/{source_to_mart['id']}/evidence",
            {
                "evidence_type": "source_field",
                "evidence_id": source_field["id"],
                "source_name": "ECIF.ecif_customer.cert_type",
                "location_text": "源字段",
                "quoted_content": "客户证件类型字段",
                "evidence_summary": "证明监管集市 cert_type 来源于 ECIF cert_type。",
            },
        )
        ybt_evidence = _post_json(
            client,
            f"{base}/mappings/mart_to_ybt/{mart_to_ybt['id']}/evidence",
            {
                "evidence_type": "db_query_result",
                "evidence_id": task_run["id"],
                "source_name": "自然语言任务结果",
                "location_text": "cert_type 枚举分布",
                "quoted_content": json.dumps(task_run["result_summary_json"], ensure_ascii=False),
                "evidence_summary": "证明 cert_type 有 01、02 和空值，需要口径说明码值转换和空值处理。",
            },
        )
        source_draft = _post_json(client, f"{base}/source-to-mart-mappings/{source_to_mart['id']}/generate-draft", {})
        ybt_draft = _post_json(client, f"{base}/mart-to-ybt-mappings/{mart_to_ybt['id']}/generate-draft", {})
        source_final = _put_json(
            client,
            f"{base}/source-to-mart-mappings/{source_to_mart['id']}",
            {
                "final_content": "业务系统到监管集市：ECIF.ecif_customer.cert_type 进入 mart_customer.cert_type，按有效客户过滤，空值列为待确认。",
                "open_questions": "请确认 ECIF 证件类型码值是否已经与最新监管代码集一致。",
            },
        )
        ybt_final = _put_json(
            client,
            f"{base}/mart-to-ybt-mappings/{mart_to_ybt['id']}",
            {
                "final_content": "监管集市到一表通：mart_customer.cert_type 映射到 CERT_TYPE，并按一表通代码集转换。",
                "open_questions": "请确认报送日期内有效客户判定规则。",
            },
        )
        legacy_source_approve = client.post(f"{base}/source-to-mart-mappings/{source_final['id']}/approve", json={"reviewed_by": "smoke"})
        legacy_ybt_approve = client.post(f"{base}/mart-to-ybt-mappings/{ybt_final['id']}/approve", json={"reviewed_by": "smoke"})
        if legacy_source_approve.status_code != 409 or legacy_ybt_approve.status_code != 409:
            raise AssertionError("治理模式双层口径旧审核接口必须返回 409")
        field_export = client.get(f"{base}/target-fields/{field['id']}/export/mapping-document", params={"format": "markdown"})
        field_export.raise_for_status()
        markdown = field_export.json()["content"]
        required_sections = [
            "一表通字段信息",
            "监管集市字段设计",
            "业务系统到监管集市取数口径",
            "监管集市到一表通取数口径",
            "参考依据",
            "待确认问题",
            "审核状态",
        ]
        missing_sections = [section for section in required_sections if section not in markdown]
        if missing_sections:
            raise AssertionError(f"导出文档缺少章节: {missing_sections}")

        excel_export = client.get(f"{base}/projects/{project_id}/export/traceability-workbook")
        excel_export.raise_for_status()
        exported_workbook = load_workbook(BytesIO(excel_export.content), data_only=True)
        exported_sheet = exported_workbook["业务口径及技术溯源"]
        export_headers = {
            str(exported_sheet.cell(row, column).value or "")
            for row in (1, 2)
            for column in range(1, exported_sheet.max_column + 1)
        }
        required_excel_headers = {
            "数据项编码", "业务口径-借记卡", "溯源-借记卡", "来源系统", "来源表英文名",
            "来源字段英文名", "处理逻辑", "技术口径确认人",
        }
        if missing := required_excel_headers - export_headers:
            raise AssertionError(f"导出 Excel 缺少表头: {sorted(missing)}")
        current_group = ""
        debit_source_columns: dict[str, int] = {}
        for column in range(1, exported_sheet.max_column + 1):
            group_value = str(exported_sheet.cell(1, column).value or "")
            if group_value:
                current_group = group_value
            child_header = str(exported_sheet.cell(2, column).value or "")
            if current_group == "溯源-借记卡" and child_header in {"来源表英文名", "来源字段英文名"}:
                debit_source_columns[child_header] = column
        exported_field_row = next(
            row for row in range(3, exported_sheet.max_row + 1)
            if str(exported_sheet.cell(row, 1).value or "") == "CERT_TYPE"
        )
        source_profile_binding = _post_json(client, f"{base}/profile-tasks/{column_profile['id']}/bind-evidence", {
            "mapping_type": "source_to_mart", "mapping_id": source_to_mart["id"],
        })
        source_profile_binding_again = _post_json(client, f"{base}/profile-tasks/{column_profile['id']}/bind-evidence", {
            "mapping_type": "source_to_mart", "mapping_id": source_to_mart["id"],
        })
        ybt_profile_binding = _post_json(client, f"{base}/profile-tasks/{column_profile['id']}/bind-evidence", {
            "mapping_type": "mart_to_ybt", "mapping_id": mart_to_ybt["id"],
        })
        if source_profile_binding_again["id"] != source_profile_binding["id"] or ybt_profile_binding["mapping_type"] != "mart_to_ybt":
            raise AssertionError("探查证据绑定两层口径未保持幂等")
        exported_source_table = exported_sheet.cell(exported_field_row, debit_source_columns["来源表英文名"]).value
        exported_source_field = exported_sheet.cell(exported_field_row, debit_source_columns["来源字段英文名"]).value
        if exported_source_table != "ecif_customer" or exported_source_field != "cert_type":
            raise AssertionError(
                f"导出 Excel 未写入已采用目录来源: {exported_source_table}.{exported_source_field}"
            )

        legacy_mapping = _post_json(
            client,
            f"{base}/fields/{field['id']}/generate-mapping",
            {
                "include_template": True,
                "include_documents": True,
                "include_sql_parse_results": True,
                "include_nl_task_results": True,
            },
        )
        draft = legacy_mapping["draft"]
        evidence_types = sorted({item["evidence_type"] for item in draft["evidences"]})

        governance_roles = ["business_analyst", "business_reviewer", "technical_analyst", "technical_reviewer", "final_reviewer"]
        governance_users = {}
        for role in governance_roles:
            password = f"smoke-only-{role}-password"
            user = _post_json(client, f"{base}/admin/users", {
                "username": f"smoke_{role}", "display_name": role, "email": f"smoke-{role}@example.invalid", "password": password,
                "institution_id": bank["id"], "institution_role": "member",
            })
            _post_json(client, f"{base}/projects/{project_id}/members", {"user_id": user["id"], "project_role": role})
            governance_users[role] = {"id": user["id"], "password": password}
        workflow_created = _post_json(client, f"{base}/projects/{project_id}/tasks/batch-create", {
            "workflow_key": "scenario_mapping_review", "targets": [{"target_field_id": field["id"], "scenario_id": debit_scenario["id"]}],
            "assignments": {role: governance_users[role]["id"] for role in governance_roles},
        })
        workflow_id = workflow_created["workflow_instance_ids"][0]
        workflow_tasks = _get_json(client, f"{base}/projects/{project_id}/tasks")
        task_by_step = {item["step_key"]: item for item in workflow_tasks if item["workflow_instance_id"] == workflow_id}
        admin_authorization = client.headers["Authorization"]
        for step, role in [("business_draft", "business_analyst"), ("business_review", "business_reviewer"), ("technical_draft", "technical_analyst"), ("technical_review", "technical_reviewer"), ("final_review", "final_reviewer")]:
            role_session = _post_json(client, f"{base}/auth/login", {"username": f"smoke_{role}", "password": governance_users[role]["password"]})
            client.headers["Authorization"] = f"Bearer {role_session['access_token']}"
            _post_json(client, f"{base}/review-tasks/{task_by_step[step]['id']}/approve", {"comment": f"{step} smoke approved"})
        final_notifications = _get_json(client, f"{base}/me/notifications")
        client.headers["Authorization"] = admin_authorization
        workflow_result = _get_json(client, f"{base}/workflows/{workflow_id}")
        confirmed_business = _get_json(client, f"{base}/scenario-business-mappings/{scenario_business['id']}")
        confirmed_lineage = _get_json(client, f"{base}/scenario-technical-lineages/{scenario_lineage['id']}")
        if confirmed_business["business_confirm_status"] != "confirmed" or confirmed_lineage["tech_confirm_status"] != "confirmed":
            raise AssertionError("五阶段最终审核未同时确认业务口径与技术溯源")
        double_workflows = _post_json(client, f"{base}/projects/{project_id}/tasks/batch-create", {
            "workflow_key": "double_layer_mapping_review",
            "targets": [
                {"target_type": "source_to_mart", "target_id": source_final["id"]},
                {"target_type": "mart_to_ybt", "target_id": ybt_final["id"]},
            ],
            "assignments": {
                "technical_reviewer": governance_users["technical_reviewer"]["id"],
                "final_reviewer": governance_users["final_reviewer"]["id"],
            },
        })
        for double_workflow_id in double_workflows["workflow_instance_ids"]:
            project_tasks = _get_json(client, f"{base}/projects/{project_id}/tasks")
            double_tasks = {item["step_key"]: item for item in project_tasks if item["workflow_instance_id"] == double_workflow_id}
            for step, role in [("technical_review", "technical_reviewer"), ("final_review", "final_reviewer")]:
                role_session = _post_json(client, f"{base}/auth/login", {"username": f"smoke_{role}", "password": governance_users[role]["password"]})
                client.headers["Authorization"] = f"Bearer {role_session['access_token']}"
                _post_json(client, f"{base}/review-tasks/{double_tasks[step]['id']}/approve", {"comment": f"{step} double layer approved"})
            client.headers["Authorization"] = admin_authorization
        source_approved = _get_json(client, f"{base}/source-to-mart-mappings/{source_final['id']}")
        ybt_approved = _get_json(client, f"{base}/mart-to-ybt-mappings/{ybt_final['id']}")
        if source_approved["mapping_status"] != "approved" or ybt_approved["mapping_status"] != "approved":
            raise AssertionError("双层口径治理工作流未完成正式审核")

        lineage_v1 = b"""-- v1 sanitized fixture
insert into mart_customer (cert_type, customer_name)
select e.cert_type, e.customer_name from ecif_customer e where e.cert_type is not null;
insert into YBT_CUSTOMER (CERT_TYPE) select cert_type from mart_customer;
"""
        lineage_v2 = b"""-- v2 sanitized fixture comment changed
insert into mart_customer (cert_type)
select case when e.cert_type='01' then 'A' else e.cert_type end
from ecif_customer e left join ecif_customer j on e.cert_type=j.cert_type
where e.cert_type <> '99';
insert into YBT_CUSTOMER (CERT_TYPE) select cert_type from mart_customer;
"""
        lineage_upload_v1 = _post_file(client, f"{base}/projects/{project_id}/scripts/upload", {"relative_path": "load_customer.sql", "dialect": "sqlite"}, {"file": ("load_customer.sql", BytesIO(lineage_v1), "text/plain")})
        shell_upload = _post_file(client, f"{base}/projects/{project_id}/scripts/upload", {"relative_path": "run_customer.sh"}, {"file": ("run_customer.sh", BytesIO(b"psql -f load_customer.sql\n"), "text/plain")})
        shell_detail = _get_json(client, f"{base}/scripts/{shell_upload['script_file_id']}")
        if not shell_detail["dependencies"] or shell_detail["dependencies"][0]["child_script_file_id"] != lineage_upload_v1["script_file_id"]:
            raise AssertionError("Shell 到 SQL 依赖未解析")
        lineage_graph = _get_json(client, f"{base}/scripts/{lineage_upload_v1['script_file_id']}/lineage?direction=both&depth=5")
        if not any(edge["transformation_expression"] for edge in lineage_graph["edges"] if edge["edge_type"] in {"derives_from", "maps_code"}):
            raise AssertionError("字段级 SQL 血缘未生成")
        lineage_excel = client.get(f"{base}/projects/{project_id}/export/lineage-workbook")
        lineage_excel.raise_for_status()
        lineage_book = load_workbook(BytesIO(lineage_excel.content))
        if "字段级血缘" not in lineage_book.sheetnames or "影响分析" not in lineage_book.sheetnames:
            raise AssertionError("血缘 Excel 缺少必需 Sheet")
        lineage_upload_v2 = _post_file(client, f"{base}/projects/{project_id}/scripts/upload", {"relative_path": "load_customer.sql", "dialect": "sqlite"}, {"file": ("load_customer.sql", BytesIO(lineage_v2), "text/plain")})
        if lineage_upload_v2["version_no"] != 2 or lineage_upload_v2["impact_severity"] != "critical":
            raise AssertionError("v2 未生成 critical 变更影响")
        lineage_impact = _get_json(client, f"{base}/lineage/impacts/{lineage_upload_v2['impact_id']}")
        if lineage_impact["workflow"]["current_step"] != "impact_analysis":
            raise AssertionError("血缘三阶段复核未自动创建")
        impact_tasks = {item["step_key"]: item for item in lineage_impact["workflow"]["tasks"]}
        for step, role in [("impact_analysis", "technical_analyst"), ("technical_review", "technical_reviewer"), ("final_review", "final_reviewer")]:
            role_session = _post_json(client, f"{base}/auth/login", {"username": f"smoke_{role}", "password": governance_users[role]["password"]})
            client.headers["Authorization"] = f"Bearer {role_session['access_token']}"
            _post_json(client, f"{base}/review-tasks/{impact_tasks[step]['id']}/approve", {"comment": f"lineage {step} smoke approved"})
        client.headers["Authorization"] = admin_authorization
        lineage_impact_reviewed = _get_json(client, f"{base}/lineage/impacts/{lineage_upload_v2['impact_id']}")
        source_after_lineage_review = _get_json(client, f"{base}/source-to-mart-mappings/{source_final['id']}")
        ybt_after_lineage_review = _get_json(client, f"{base}/mart-to-ybt-mappings/{ybt_final['id']}")
        if lineage_impact_reviewed["status"] != "reviewed" or source_after_lineage_review["lineage_status"] != "verified" or ybt_after_lineage_review["lineage_status"] != "verified":
            raise AssertionError("血缘最终审核后未恢复 verified")
        batch_job_response = client.post(f"{base}/projects/{project_id}/batch/generate-business-drafts", headers={"Idempotency-Key": "smoke-business-batch"}, json={"field_ids": []})
        batch_job_response.raise_for_status()
        batch_job = _get_json(client, f"{base}/jobs/{batch_job_response.json()['id']}")
        audit_rows = _get_json(client, f"{base}/audit?project_id={project_id}")
        other_bank = _post_json(client, f"{base}/admin/institutions", {"institution_code": "OTHER_BANK", "institution_name": "其他示例银行", "institution_type": "bank"})
        outsider_password = "smoke-only-outsider-password"
        _post_json(client, f"{base}/admin/users", {"username": "smoke_outsider", "display_name": "跨行用户", "email": "smoke-outsider@example.invalid", "password": outsider_password, "institution_id": other_bank["id"], "institution_role": "member"})
        outsider_session = _post_json(client, f"{base}/auth/login", {"username": "smoke_outsider", "password": outsider_password})
        cross_bank_response = client.get(f"{base}/projects/{project_id}", headers={"Authorization": f"Bearer {outsider_session['access_token']}"})
        if cross_bank_response.status_code != 404:
            raise AssertionError(f"Cross-bank project read must be hidden, got {cross_bank_response.status_code}")
        cross_bank_lineage = client.get(f"{base}/scripts/{lineage_upload_v1['script_file_id']}", headers={"Authorization": f"Bearer {outsider_session['access_token']}"})
        if cross_bank_lineage.status_code != 404:
            raise AssertionError(f"Cross-bank lineage read must be hidden, got {cross_bank_lineage.status_code}")
        client.headers["Authorization"] = admin_authorization

        delivery_smoke = _run_delivery_smoke(
            client, base, project_id, debit_scenario, mart_field,
            governance_users, outsider_session["access_token"], admin_authorization,
        )

        output = {
            "project_id": project_id,
            "knowledge_document_ids": [item["id"] for item in knowledge_documents],
            "knowledge_unit_count": len(knowledge_units),
            "knowledge_search_types": sorted(search_types),
            "knowledge_citation_count": len(grounded["citations"]),
            "evaluation_case_id": evaluation_case["id"],
            "evaluation_run_id": evaluation_run["id"],
            "evaluation_metrics": evaluation_run["summary_metrics_json"],
            "template_status": template["parse_status"],
            "template_field_count": template["field_count"],
            "apply_result": apply_result,
            "field_id": field["id"],
            "traceability_template_status": traceability_template["parse_status"],
            "traceability_preview_rows": len(traceability_preview["results"][0]["parsed_rows_json"]),
            "traceability_apply": traceability_apply,
            "scenario_codes": [item["scenario_code"] for item in scenarios],
            "recommendation_top_score": source_recommendations["recommendations"][0]["score"],
            "metadata_sync_status": metadata_sync["status"],
            "catalog_table_count": catalog_tables["total"],
            "catalog_source_field_id": catalog_source_import["source_field_id"],
            "column_profile_status": column_profile["status"],
            "column_profile_null_rate": column_profile["profile_result_json"].get("null_rate"),
            "sensitive_profile_top_values": sensitive_profile["profile_result_json"].get("top_values"),
            "selected_recommendation_id": selected_recommendation["recommendation"]["id"],
            "scenario_business_confirm_status": confirmed_business["business_confirm_status"],
            "scenario_technical_confirm_status": confirmed_lineage["tech_confirm_status"],
            "adopted_business_chars": len(adopted_business["final_content"]),
            "adopted_lineage_chars": len(adopted_lineage["final_content"]),
            "export_excel_bytes": len(excel_export.content),
            "datasource_id": datasource["id"],
            "datasource_test": datasource_test,
            "task_status": task_run["status"],
            "extracted": {
                "datasource": task_run["datasource_name"],
                "table": task_run["extracted_table_name"],
                "field": task_run["extracted_field_name"],
            },
            "null_profile": task_run["result_summary_json"]["null_profile"],
            "distinct_profile": task_run["result_summary_json"]["distinct_profile"],
            "enum_distribution_rows": task_run["result_summary_json"]["enum_distribution"]["row_count"],
            "business_system_id": business_system["id"],
            "source_field_id": source_field["id"],
            "mart_field_id": mart_field["id"],
            "source_to_mart_mapping_status": source_approved["mapping_status"],
            "mart_to_ybt_mapping_status": ybt_approved["mapping_status"],
            "legacy_double_layer_review_statuses": [legacy_source_approve.status_code, legacy_ybt_approve.status_code],
            "source_draft_confidence": source_draft["confidence_level"],
            "ybt_draft_confidence": ybt_draft["confidence_level"],
            "mapping_evidence_ids": [source_evidence["id"], ybt_evidence["id"], source_profile_binding["id"], ybt_profile_binding["id"]],
            "export_markdown_chars": len(markdown),
            "draft_id": draft["id"],
            "evidence_types": evidence_types,
            "template_reference_summary": draft.get("template_reference_summary"),
            "db_query_summary": draft.get("db_query_summary"),
            "evidence_completeness": draft.get("evidence_completeness"),
            "governance_institution_id": bank["id"],
            "workflow_status": workflow_result["status"],
            "workflow_decision_count": len(workflow_result["decisions"]),
            "notification_count": len(final_notifications),
            "background_job_status": batch_job["status"],
            "audit_log_count": len(audit_rows),
            "cross_bank_status": cross_bank_response.status_code,
            "lineage_script_version": lineage_upload_v2["version_no"],
            "lineage_edge_count": len(lineage_graph["edges"]),
            "lineage_change_categories": lineage_upload_v2["change_categories"],
            "lineage_impact_status": lineage_impact_reviewed["status"],
            "lineage_excel_bytes": len(lineage_excel.content),
            "cross_bank_lineage_status": cross_bank_lineage.status_code,
            "deliverable": delivery_smoke,
        }
        print(json.dumps(output, ensure_ascii=False, indent=2))


def _run_delivery_smoke(client, base, project_id, scenario, mart_field, users, outsider_token, admin_authorization):
    table = _post_json(client, f"{base}/target-tables", {"project_id": project_id, "table_code": "YBT_DELIVERY", "table_name": "正式交付验收表"})
    field = _post_json(client, f"{base}/fields", {"project_id": project_id, "target_table_id": table["id"], "field_code": "DELIVERY_CERT_TYPE", "field_name": "交付证件类型", "field_type": "VARCHAR(20)", "regulatory_description": "正式交付验收使用的脱敏监管定义"})
    business = _post_json(client, f"{base}/target-fields/{field['id']}/scenarios/{scenario['id']}/business-mapping", {"business_definition": "借记卡有效客户证件类型", "final_content": "仅包含当前有效借记卡客户，排除注销客户。"})
    technical = _post_json(client, f"{base}/target-fields/{field['id']}/scenarios/{scenario['id']}/technical-lineage", {"business_mapping_id": business["id"], "source_system_name": "ECIF", "source_database_name": "ECIF_DB", "source_schema_name": "main", "source_table_english_name": "ecif_customer", "source_field_english_name": "cert_type", "processing_logic": "直接取值并按监管代码映射", "final_content": "ECIF_DB.main.ecif_customer.cert_type"})
    _post_json(client, f"{base}/mappings/scenario_business/{business['id']}/evidence", {"evidence_type": "manual_note", "source_name": "脱敏业务确认记录", "evidence_summary": "业务范围已由业务人员核验"})
    _post_json(client, f"{base}/mappings/scenario_technical/{technical['id']}/evidence", {"evidence_type": "catalog_column", "source_name": "ecif_customer.cert_type", "evidence_summary": "物理来源字段已由科技人员核验"})
    workflow = _post_json(client, f"{base}/projects/{project_id}/tasks/batch-create", {"workflow_key": "scenario_mapping_review", "targets": [{"target_field_id": field["id"], "scenario_id": scenario["id"]}], "assignments": {role: users[role]["id"] for role in users}})
    workflow_id = workflow["workflow_instance_ids"][0]
    tasks = {item["step_key"]: item for item in _get_json(client, f"{base}/projects/{project_id}/tasks") if item["workflow_instance_id"] == workflow_id}
    for step, role in [("business_draft", "business_analyst"), ("business_review", "business_reviewer"), ("technical_draft", "technical_analyst"), ("technical_review", "technical_reviewer"), ("final_review", "final_reviewer")]:
        session = _post_json(client, f"{base}/auth/login", {"username": f"smoke_{role}", "password": users[role]["password"]})
        client.headers["Authorization"] = f"Bearer {session['access_token']}"
        _post_json(client, f"{base}/review-tasks/{tasks[step]['id']}/approve", {"comment": f"deliverable {step} approved"})
    client.headers["Authorization"] = admin_authorization
    mart_mapping = _post_json(client, f"{base}/target-fields/{field['id']}/mart-to-ybt-mappings", {"mart_field_id": mart_field["id"], "mapping_name": "正式交付集市映射", "final_content": "mart_customer.cert_type 映射至 DELIVERY_CERT_TYPE"})
    _post_json(client, f"{base}/mappings/mart_to_ybt/{mart_mapping['id']}/evidence", {"evidence_type": "manual_note", "source_name": "脱敏集市映射确认", "evidence_summary": "集市到一表通映射已由科技人员核验"})
    double = _post_json(client, f"{base}/projects/{project_id}/tasks/batch-create", {"workflow_key": "double_layer_mapping_review", "targets": [{"target_type": "mart_to_ybt", "target_id": mart_mapping["id"]}], "assignments": {"technical_reviewer": users["technical_reviewer"]["id"], "final_reviewer": users["final_reviewer"]["id"]}})
    double_id = double["workflow_instance_ids"][0]
    double_tasks = {item["step_key"]: item for item in _get_json(client, f"{base}/projects/{project_id}/tasks") if item["workflow_instance_id"] == double_id}
    for step, role in [("technical_review", "technical_reviewer"), ("final_review", "final_reviewer")]:
        session = _post_json(client, f"{base}/auth/login", {"username": f"smoke_{role}", "password": users[role]["password"]})
        client.headers["Authorization"] = f"Bearer {session['access_token']}"
        _post_json(client, f"{base}/review-tasks/{double_tasks[step]['id']}/approve", {"comment": f"deliverable {step} approved"})
    client.headers["Authorization"] = admin_authorization

    template_content = _delivery_template_bytes()
    template = _post_file(client, f"{base}/projects/{project_id}/deliverable-templates/upload", {"template_name": "银行正式交付模板", "template_type": "full_delivery_package"}, {"file": ("银行正式交付模板.xlsx", BytesIO(template_content), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")})
    version_id = template["version"]["id"]
    sheet_mappings = [
        {"business_section": "target_field", "sheet_name": "业务口径及技术溯源表", "header_row_start": 2, "header_row_end": 2, "data_start_row": 3, "columns": [{"business_field": "target_field_code", "excel_column": "A", "required": True}, {"business_field": "target_field_name", "excel_column": "B", "required": True}, {"business_field": "regulatory_definition", "excel_column": "C"}]},
        {"business_section": "source_to_mart", "sheet_name": "业务系统到监管集市口径", "data_start_row": 2, "columns": [{"business_field": "source_to_mart_final_content", "excel_column": "A"}]},
        {"business_section": "mart_to_ybt", "sheet_name": "监管集市到一表通口径", "data_start_row": 2, "columns": [{"business_field": "mart_to_ybt_final_content", "excel_column": "A"}]},
        {"business_section": "evidence", "sheet_name": "证据引用清单", "data_start_row": 2, "columns": [{"business_field": "evidence_type", "excel_column": "A"}, {"business_field": "evidence_summary", "excel_column": "B"}]},
        {"business_section": "pending_question", "sheet_name": "待确认问题", "data_start_row": 2, "columns": [{"business_field": "question_text", "excel_column": "A"}, {"business_field": "question_status", "excel_column": "B"}]},
        {"business_section": "review_record", "sheet_name": "审核记录", "data_start_row": 2, "columns": [{"business_field": "action", "excel_column": "A"}, {"business_field": "approved_at", "excel_column": "B"}]},
    ]
    _post_json(client, f"{base}/deliverable-template-versions/{version_id}/configure", {"sheet_mappings": sheet_mappings})
    preview = client.post(f"{base}/deliverable-template-versions/{version_id}/preview-render", json={})
    preview.raise_for_status(); load_workbook(BytesIO(preview.content), data_only=False)

    historical = _post_file(client, f"{base}/projects/{project_id}/historical-calibers/upload", {"import_name": "历史业务和技术口径", "document_type": "full_package"}, {"file": ("历史口径.xlsx", BytesIO(_delivery_history_bytes()), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")})
    historical_item = historical["items"][0]
    if historical_item["match_status"] != "matched": raise AssertionError("历史口径未自动唯一匹配")
    reuse = _post_json(client, f"{base}/historical-caliber-items/{historical_item['id']}/reuse", {})
    if reuse["final_content_overwritten"]: raise AssertionError("历史复用覆盖了人工 final_content")

    question = _post_json(client, f"{base}/projects/{project_id}/questions", {"target_table_id": table["id"], "target_field_id": field["id"], "scenario_id": scenario["id"], "question_type": "code_mapping", "question_text": "确认码值映射版本", "priority": "high", "assigned_role": "technical_analyst"})
    package = _post_json(client, f"{base}/projects/{project_id}/deliverables", {"package_name": "正式交付验收包", "target_table_id": table["id"], "template_version_id": version_id})
    _post_json(client, f"{base}/deliverables/{package['id']}/generate", {})
    blocked = _post_json(client, f"{base}/deliverables/{package['id']}/validate", {})
    if not any(item["code"] == "high_priority_question" for item in blocked["issues"]): raise AssertionError("高优问题未阻止正式批准")
    _post_json(client, f"{base}/questions/{question['id']}/answer", {"resolution_text": "采用当前监管代码集"}); _post_json(client, f"{base}/questions/{question['id']}/accept", {})
    validation = _post_json(client, f"{base}/deliverables/{package['id']}/validate", {})
    if validation["error_count"]: raise AssertionError(f"正式交付校验仍有错误: {validation['issues']}")
    rendered = _post_json(client, f"{base}/deliverables/{package['id']}/render", {})
    downloaded = client.get(f"{base}/deliverables/{package['id']}/download"); downloaded.raise_for_status()
    workbook = load_workbook(BytesIO(downloaded.content), data_only=False); sheet = workbook["业务口径及技术溯源表"]
    if sheet["A3"].value != "DELIVERY_CERT_TYPE" or sheet["D3"].value != "=1+1" or "A1:C1" not in [str(item) for item in sheet.merged_cells.ranges] or not sheet["A2"].font.bold: raise AssertionError("正式 Excel 未保留内容、公式、样式或合并单元格")
    approved1 = _post_json(client, f"{base}/deliverables/{package['id']}/approve", {})
    _post_json(client, f"{base}/deliverables/{package['id']}/render", {})
    approved2 = _post_json(client, f"{base}/deliverables/{package['id']}/approve", {})
    comparison = _post_json(client, f"{base}/projects/{project_id}/caliber-comparisons", {"left_package_version_id": approved1["version"]["id"], "right_package_version_id": approved2["version"]["id"], "left": {"source_field": "cert_type"}, "right": {"source_field": "cert_type_v2"}})
    cross = client.get(f"{base}/deliverables/{package['id']}", headers={"Authorization": f"Bearer {outsider_token}"})
    if cross.status_code != 404: raise AssertionError("跨银行用户可读取正式交付包")
    client.headers["Authorization"] = admin_authorization
    return {"template_version_id": version_id, "historical_item_id": historical_item["id"], "package_id": package["id"], "rendered_file_id": rendered["file_id"], "validation_errors": validation["error_count"], "approved_versions": [approved1["version"]["version_no"], approved2["version"]["version_no"]], "comparison_id": comparison["id"], "cross_bank_status": cross.status_code, "sheet_count": len(workbook.sheetnames)}


def _delivery_template_bytes():
    workbook = Workbook(); sheet = workbook.active; sheet.title = "业务口径及技术溯源表"; sheet.merge_cells("A1:C1"); sheet["A1"] = "银行正式交付模板"; sheet["A2"] = "字段代码"; sheet["B2"] = "字段名称"; sheet["C2"] = "监管定义"; sheet["D3"] = "=1+1"; sheet["A2"].font = Font(bold=True); sheet.freeze_panes = "A3"
    for name in ["业务系统到监管集市口径", "监管集市到一表通口径", "字段级血缘", "证据引用清单", "待确认问题", "历史口径差异", "脚本变更影响", "审核记录", "版本说明"]:
        target = workbook.create_sheet(name); target.append([name, "状态", "说明"])
    stream = BytesIO(); workbook.save(stream); return stream.getvalue()


def _delivery_history_bytes():
    workbook = Workbook(); sheet = workbook.active; sheet.title = "历史业务口径"; sheet.append(["字段代码", "字段名称", "业务场景", "业务口径", "技术溯源", "来源系统", "来源schema", "来源表英文名", "来源字段英文名"]); sheet.append(["DELIVERY_CERT_TYPE", "交付证件类型", "借记卡", "历史有效客户业务口径", "历史 ECIF 技术溯源", "ECIF", "main", "ecif_customer", "cert_type"]); stream = BytesIO(); workbook.save(stream); return stream.getvalue()


def _upload_knowledge(
    client: httpx.Client,
    base: str,
    project_id: int,
    path: Path,
    knowledge_type: str,
    knowledge_scope: str = "project",
    institution_name: str = "",
) -> dict:
    with path.open("rb") as file:
        return _post_file(
            client,
            f"{base}/projects/{project_id}/knowledge/documents/upload",
            data={
                "knowledge_type": knowledge_type,
                "knowledge_scope": knowledge_scope,
                "institution_name": institution_name,
                "confidentiality_level": "internal",
            },
            files={"file": (path.name, file, "application/octet-stream")},
        )


def _write_qa(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "答疑"
    sheet.append(["问题", "监管回复", "表代码", "字段代码", "字段名称", "业务场景"])
    sheet.append([
        "借记卡客户证件类型应如何取值",
        "优先采用 ECIF_CUSTOMER.CERT_TYPE，并按一表通代码集转换。",
        "YBT_CUSTOMER",
        "CERT_TYPE",
        "客户证件类型",
        "借记卡",
    ])
    workbook.save(path)


def _write_historical_mapping(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "历史口径"
    sheet.append(["字段代码", "字段名称", "来源系统", "来源表", "来源字段", "处理逻辑"])
    sheet.append(["CERT_TYPE", "客户证件类型", "ECIF", "ecif_customer", "cert_type", "有效客户范围内直接取值并转换码值"])
    workbook.save(path)


def _write_policy_docx(path: Path) -> None:
    document = Document()
    document.add_heading("客户证件类型监管范围", level=1)
    document.add_paragraph("借记卡客户证件类型应有真实来源字段和可验证数据质量证据。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "CERT_TYPE"
    document.save(path)


def _write_text_pdf(path: Path) -> None:
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length 62 >>\nstream\nBT /F1 12 Tf 72 720 Td (Regulatory CERT_TYPE guidance) Tj ET\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, value in enumerate(objects, 1):
        offsets.append(len(content))
        content += f"{number} 0 obj\n".encode() + value + b"\nendobj\n"
    xref_offset = len(content)
    content += f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode()
    content += b"".join(f"{offset:010d} 00000 n \n".encode() for offset in offsets[1:])
    content += f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n".encode()
    path.write_bytes(content)


def _write_template(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "客户信息表"
    sheet.append(["表编号", "表名称", "字段代码", "字段中文名", "字段类型", "是否必填", "字段定义", "监管说明"])
    sheet.append(["YBT_CUSTOMER", "客户信息表", "CERT_TYPE", "客户证件类型", "varchar(20)", "是", "客户身份证件类型", "按监管代码集转换"])
    sheet.append(["YBT_CUSTOMER", "客户信息表", "CUSTOMER_ID", "客户编号", "varchar(64)", "是", "客户唯一编号", "不能为空"])
    workbook.save(path)


def _write_traceability_template(path: Path) -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "客户信息场景口径"
    sheet["A1"] = "脱敏模拟业务口径及技术溯源表"
    fixed_headers = [
        "数据项编码", "数据项名称", "数据类别", "数据格式", "字段业务定义（监管原始口径）",
        "字段业务定义（监管定义细化）", "报表名称", "字段名称", "EAST口径", "字段业务定义（行内）", "备注",
    ]
    for column, title in enumerate(fixed_headers, start=1):
        sheet.cell(2, column, title)
        sheet.merge_cells(start_row=2, start_column=column, end_row=3, end_column=column)
    business_headers = ["字段业务定义", "源系统截图", "源系统改造", "外部数据", "手工补录", "业务口径确认人", "备注"]
    technical_headers = [
        "来源系统", "来源库", "来源schema", "来源表英文名", "来源表中文名", "来源字段英文名",
        "来源字段中文名", "处理逻辑", "处理逻辑类型", "技术口径确认人", "备注",
    ]
    column = 12
    for scenario_name, source_system in [("借记卡", "借记卡系统"), ("信用卡", "信用卡系统"), ("贷款产品", "信贷系统")]:
        column = _write_group(sheet, column, f"业务口径-{scenario_name}", business_headers)
        column = _write_group(sheet, column, f"溯源-{scenario_name}", technical_headers)
    rows = [
        ["CERT_TYPE", "客户证件类型", "基础信息", "VARCHAR(20)", "客户身份证件类型", "按产品场景确认客户证件类型", "客户信息表", "客户证件类型", "EAST_CERT_TYPE", "行内证件类型", "脱敏模拟"],
        ["CUSTOMER_ID", "客户编号", "基础信息", "VARCHAR(64)", "客户唯一编号", "按产品场景确认客户范围", "客户信息表", "客户编号", "EAST_CUSTOMER_ID", "行内客户编号", "脱敏模拟"],
    ]
    for row_number, fixed_values in enumerate(rows, start=4):
        for fixed_column, value in enumerate(fixed_values, start=1):
            sheet.cell(row_number, fixed_column, value)
        group_column = 12
        for scenario_name, source_system in [("借记卡", "借记卡系统"), ("信用卡", "信用卡系统"), ("贷款产品", "信贷系统")]:
            field_code = fixed_values[0]
            _write_values(sheet, row_number, group_column, [f"{scenario_name}{fixed_values[1]}业务口径", "是", "否", "否", "否", f"{scenario_name}业务部门", "待确认"])
            group_column += len(business_headers)
            _write_values(sheet, row_number, group_column, [source_system, "DEMO_DB", "ODS", f"{scenario_name}_CUSTOMER", f"{scenario_name}客户表", field_code, fixed_values[1], "源字段直接取值", "direct", "信息科技部", "脱敏模拟"])
            group_column += len(technical_headers)
    workbook.save(path)


def _write_group(sheet, start_column: int, title: str, headers: list[str]) -> int:
    end_column = start_column + len(headers) - 1
    sheet.cell(2, start_column, title)
    sheet.merge_cells(start_row=2, start_column=start_column, end_row=2, end_column=end_column)
    for offset, header in enumerate(headers):
        sheet.cell(3, start_column + offset, header)
    return end_column + 1


def _write_values(sheet, row: int, start_column: int, values: list[str]) -> None:
    for offset, value in enumerate(values):
        sheet.cell(row, start_column + offset, value)


def _write_sqlite_source(path: Path) -> None:
    connection = sqlite3.connect(path)
    try:
        connection.execute("create table ecif_customer (cert_type text, customer_name text)")
        connection.executemany(
            "insert into ecif_customer (cert_type, customer_name) values (?, ?)",
            [("01", "张三"), ("01", "李四"), ("02", "王五"), (None, "赵六")],
        )
        connection.commit()
    finally:
        connection.close()


def _post_json(client: httpx.Client, url: str, payload: dict) -> dict:
    response = client.post(url, json=payload)
    response.raise_for_status()
    return response.json()


def _get_json(client: httpx.Client, url: str) -> dict | list[dict]:
    response = client.get(url)
    response.raise_for_status()
    return response.json()


def _post_file(client: httpx.Client, url: str, data: dict, files: dict) -> dict:
    response = client.post(url, data=data, files=files)
    response.raise_for_status()
    return response.json()


def _put_json(client: httpx.Client, url: str, payload: dict) -> dict:
    response = client.put(url, json=payload)
    response.raise_for_status()
    return response.json()


if __name__ == "__main__":
    main()
