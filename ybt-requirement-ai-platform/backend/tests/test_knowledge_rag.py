from contextlib import contextmanager
from io import BytesIO
from pathlib import Path
from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook
from pypdf import PdfWriter
import pytest
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker
from sqlalchemy.pool import StaticPool
from app.core.database import Base,get_db
from app.main import app
from app.models import ModelCallLog,ModelProfile,Project,PromptTemplateVersion,RetrievalLog
from app.services.db.dialect import qualify_table,quote_identifier
from app.services.embeddings.factory import get_embedding_service
from app.services.embeddings.openai_compatible import OpenAICompatibleEmbeddingService
from app.services.rag.citation_validator import validate_citations
from app.services.llm.prompt_runtime import get_prompt_runtime,prepare_model_input,record_model_call
from app.services.security import ensure_external_allowed, redact_content
from app.services.vector import VectorRecord
from app.services.vector.factory import get_vector_store
from app.services.vector.milvus import MilvusVectorStore

def test_knowledge_versions_hybrid_search_grounded_answer_and_evaluation(tmp_path:Path,monkeypatch):
    monkeypatch.setenv("STORAGE_DIR",str(tmp_path));get_vector_store.cache_clear();get_embedding_service.cache_clear()
    with _client() as client:
        project=_post(client,"/api/projects",{"name":"知识项目","bank_name":"甲银行"})
        payload=_qa_excel("监管答疑：客户证件类型应来自 ECIF_CUSTOMER.CERT_TYPE")
        document=_upload(client,project["id"],"监管答疑.xlsx",payload,"regulatory_qa","institution","甲银行")
        repeated=_upload(client,project["id"],"监管答疑.xlsx",payload,"regulatory_qa","institution","甲银行")
        assert repeated["id"]==document["id"];assert len(_get(client,f"/api/knowledge/documents/{document['id']}/versions"))==1
        units=_get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={document['id']}");assert len(units)==1;unit=units[0];assert unit["unit_type"]=="qa";assert unit["source_sheet_name"]=="答疑";assert unit["source_cell_range"]=="A2:F2"
        search=_post(client,f"/api/projects/{project['id']}/knowledge/hybrid-search",{"query":"客户证件类型 CERT_TYPE","knowledge_types":["regulatory_qa"],"top_k":5});assert search["items"][0]["knowledge_unit_id"]==unit["id"];assert search["items"][0]["source_cell_range"]=="A2:F2"
        answer=_post(client,f"/api/projects/{project['id']}/knowledge/ask",{"query":"客户证件类型取哪个字段","top_k":5});assert answer["citations"][0]["knowledge_unit_id"]==unit["id"]
        case=_post(client,f"/api/projects/{project['id']}/evaluations/cases",{"case_name":"证件类型召回","query_text":"客户证件类型 CERT_TYPE","expected_knowledge_unit_ids_json":[unit["id"]],"expected_answer_keywords_json":["CERT_TYPE"]});run=_post(client,f"/api/projects/{project['id']}/evaluations/runs",{"run_name":"回归"});assert run["summary_metrics_json"]["recall_at_5"]==1;assert run["summary_metrics_json"]["mrr"]==1
        feedback=_post(client,f"/api/projects/{project['id']}/feedback",{"feedback_type":"retrieval","target_type":"knowledge_unit","target_id":unit["id"],"rating":"correct","comment":"引用准确"});assert feedback["rating"]=="correct"
        changed=_upload(client,project["id"],"监管答疑.xlsx",_qa_excel("更新答疑：优先取 ECIF_CUSTOMER.CERT_TYPE"),"regulatory_qa","institution","甲银行");assert changed["id"]==document["id"];assert len(_get(client,f"/api/knowledge/documents/{document['id']}/versions"))==2;assert len(_get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={document['id']}"))==1
        other=_post(client,"/api/projects",{"name":"隔离项目","bank_name":"乙银行"});isolated=_post(client,f"/api/projects/{other['id']}/knowledge/hybrid-search",{"query":"CERT_TYPE","top_k":5});assert isolated["items"]==[]

def test_docx_pdf_text_markdown_and_sql_ingestion_preserve_locations(tmp_path:Path,monkeypatch):
    monkeypatch.setenv("STORAGE_DIR",str(tmp_path));get_vector_store.cache_clear();get_embedding_service.cache_clear()
    with _client() as client:
        project=_post(client,"/api/projects",{"name":"多格式知识"})
        document=Document();document.add_heading("第一章 监管范围",level=1);document.add_paragraph("本条规定客户证件类型的适用范围。");table=document.add_table(rows=1,cols=2);table.cell(0,0).text="字段";table.cell(0,1).text="CERT_TYPE";stream=BytesIO();document.save(stream)
        docx=_upload(client,project["id"],"制度.docx",stream.getvalue(),"regulatory_policy");docx_units=_get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={docx['id']}");assert any(item["source_heading"]=="第一章 监管范围" for item in docx_units);assert any(item["source_cell_range"]=="table:1:row:1" for item in docx_units)
        writer=PdfWriter();writer.add_blank_page(width=200,height=200);pdf=BytesIO();writer.write(pdf);pdf_doc=_upload(client,project["id"],"说明.pdf",pdf.getvalue(),"regulatory_policy");assert "第 1 页" in str(pdf_doc["warnings_json"])
        txt=_upload(client,project["id"],"调研.txt","业务调研结论：借记卡客户范围待确认。".encode(),"business_research");assert _get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={txt['id']}")[0]["unit_type"]=="paragraph"
        md=_upload(client,project["id"],"解释.md","## 字段解释\n\n客户证件类型代码。".encode(),"field_explanation");assert _get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={md['id']}")
        sql=_upload(client,project["id"],"历史.sql",b"select cert_type, case when cert_type='01' then 'ID' end from ecif_customer where status='A'","sql_evidence");sql_unit=_get(client,f"/api/projects/{project['id']}/knowledge/units?document_id={sql['id']}")[0];assert "ecif_customer" in sql_unit["content"];assert "cert_type" in sql_unit["content"]

def test_dialect_identifiers_and_datasource_guard():
    assert quote_identifier("order","mysql")=="`order`";assert qualify_table("ODS","CUSTOMER","mysql_compatible")=="`ODS`.`CUSTOMER`";assert qualify_table("main","customer","sqlite")=='"customer"'


def test_scope_reuse_isolation_and_no_evidence_answer(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path))
    get_vector_store.cache_clear()
    get_embedding_service.cache_clear()
    with _client() as client:
        owner = _post(client, "/api/projects", {"name": "甲行项目一", "bank_name": "甲银行"})
        same_bank = _post(client, "/api/projects", {"name": "甲行项目二", "bank_name": "甲银行"})
        other_bank = _post(client, "/api/projects", {"name": "乙行项目", "bank_name": "乙银行"})
        _upload(client, owner["id"], "通用制度.txt", "监管通用证件类型规则".encode(), "regulatory_policy", "global")
        _upload(client, owner["id"], "甲行口径.txt", "甲银行内部证件类型规则".encode(), "historical_mapping", "institution", "甲银行")

        same_result = _post(client, f"/api/projects/{same_bank['id']}/knowledge/hybrid-search", {"query": "证件类型规则", "top_k": 10})
        assert {item["knowledge_type"] for item in same_result["items"]} == {"regulatory_policy", "historical_mapping"}
        assert all(item["vector_score"] > 0 for item in same_result["items"])

        other_result = _post(client, f"/api/projects/{other_bank['id']}/knowledge/hybrid-search", {"query": "证件类型规则", "top_k": 10})
        assert {item["knowledge_type"] for item in other_result["items"]} == {"regulatory_policy"}
        no_evidence = _post(client, f"/api/projects/{other_bank['id']}/knowledge/ask", {"query": "完全不存在的量子字段 ZXQ_999", "knowledge_types": ["sql_evidence"], "top_k": 5})
        assert no_evidence["citations"] == []
        assert "待确认" in no_evidence["answer"]


def test_redaction_external_policy_and_citation_validation(db_session):
    raw = "手机号 13800138000，证件号 110101199001011234，邮箱 user@example.com，password=secret"
    redacted = redact_content(raw)
    assert "13800138000" not in redacted
    assert "110101199001011234" not in redacted
    assert "user@example.com" not in redacted
    assert "secret" not in redacted
    with pytest.raises(ValueError, match="restricted"):
        ensure_external_allowed("restricted", local_only=False)
    with pytest.raises(ValueError, match="confidential"):
        ensure_external_allowed("confidential", local_only=False)
    validate_citations(db_session, [])
    with pytest.raises(ValueError, match="不存在"):
        validate_citations(db_session, [{"knowledge_unit_id": 999999}])


def test_openai_compatible_embedding_and_milvus_adapter(monkeypatch):
    class Response:
        def raise_for_status(self):
            return None

        def json(self):
            return {"data": [{"embedding": [0.1, 0.2]}, {"embedding": [0.3, 0.4]}]}

    calls = []

    def fake_post(url, **kwargs):
        calls.append((url, kwargs))
        return Response()

    monkeypatch.setenv("TEST_EMBEDDING_KEY", "not-a-real-secret")
    monkeypatch.setattr("app.services.embeddings.openai_compatible.httpx.post", fake_post)
    embedding = OpenAICompatibleEmbeddingService("http://embedding.test/v1", "demo-model", "TEST_EMBEDDING_KEY")
    assert embedding.embed_texts(["a", "b"]) == [[0.1, 0.2], [0.3, 0.4]]
    assert calls[0][0] == "http://embedding.test/v1/embeddings"
    assert calls[0][1]["json"] == {"model": "demo-model", "input": ["a", "b"]}

    class FakeMilvusClient:
        def __init__(self):
            self.created = []
            self.upserts = []
            self.searches = []
            self.deletes = []

        def has_collection(self, name):
            return False

        def create_collection(self, **kwargs):
            self.created.append(kwargs)

        def upsert(self, collection_name, data):
            self.upserts.append((collection_name, data))

        def search(self, collection_name, vectors, **kwargs):
            self.searches.append((collection_name, vectors, kwargs))
            return [[{"id": "knowledge-unit-1", "distance": 0.9, "entity": {"content": "证据", "project_id": 1}}]]

        def delete(self, collection_name, **kwargs):
            self.deletes.append((collection_name, kwargs))

    client = FakeMilvusClient()
    store = MilvusVectorStore(client=client)
    store.upsert([VectorRecord("knowledge-unit-1", [0.1, 0.2], "证据", {"project_id": 1})])
    result = store.search([0.1, 0.2], 5, {"project_id": 1, "knowledge_type": ["regulatory_qa"]})
    store.delete(filters={"project_id": 1})
    assert result[0].metadata["project_id"] == 1
    assert client.created[0]["dimension"] == 2
    assert 'knowledge_type in ["regulatory_qa"]' in client.searches[0][2]["filter"]
    assert client.deletes[0][1]["filter"] == "project_id == 1"


def test_model_profile_rejects_nested_credentials(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("STORAGE_DIR", str(tmp_path))
    with _client() as client:
        response = client.post("/api/model-profiles", json={
            "profile_name": "unsafe-profile",
            "provider_type": "openai_compatible",
            "config_json": {"headers": {"AuthorizationToken": "plaintext-secret"}},
        })
        assert response.status_code == 400
        assert "credentials" in response.text


def test_prompt_version_external_policy_and_model_call_audit(db_session):
    project=Project(name="Prompt 审计项目");db_session.add(project);db_session.flush()
    profile=ModelProfile(profile_name="external",provider_type="openai_compatible",base_url="https://model.example/v1",model_name="demo",enabled=True,local_only=False,config_json={"api_key_env_name":"TEST_MODEL_KEY"});db_session.add(profile)
    prompt=PromptTemplateVersion(prompt_key="scenario_business_mapping",version_no=7,system_prompt="场景业务口径",user_prompt_template="{evidence}",enabled=True);db_session.add(prompt)
    retrieval=RetrievalLog(project_id=project.id,query_text="证件类型",query_type="hybrid",filters_json={},retrieval_strategy="test",result_ids_json=[]);db_session.add(retrieval);db_session.commit()
    runtime=get_prompt_runtime(db_session,"scenario_business_mapping")
    assert runtime.version==7 and runtime.model_profile_id==profile.id and runtime.local_only is False
    with pytest.raises(ValueError,match="restricted"):
        prepare_model_input(runtime,"受限知识",["restricted"])
    model_input=prepare_model_input(runtime,"联系电话 13800138000",["internal"])
    assert "13800138000" not in model_input
    record_model_call(db_session,project.id,runtime,model_input,{"draft":"联系电话 13800138000"},retrieval_log_id=retrieval.id)
    db_session.commit();log=db_session.query(ModelCallLog).one()
    assert log.prompt_version==7 and log.retrieval_log_id==retrieval.id
    assert "13800138000" not in (log.output_summary or "")

def _qa_excel(answer):
    workbook=Workbook();sheet=workbook.active;sheet.title="答疑";sheet.append(["问题","监管回复","表代码","字段代码","字段名称","备注"]);sheet.append(["客户证件类型如何取值",answer,"YBT_CUSTOMER","CERT_TYPE","客户证件类型","脱敏模拟"]);stream=BytesIO();workbook.save(stream);return stream.getvalue()
def _upload(client,project_id,name,content,kind,scope="project",institution=None):
    response=client.post(f"/api/projects/{project_id}/knowledge/documents/upload",data={"knowledge_type":kind,"knowledge_scope":scope,"institution_name":institution or "","confidentiality_level":"internal"},files={"file":(name,content,"application/octet-stream")});assert response.status_code==200,response.text;return response.json()
def _post(client,path,payload):
    response=client.post(path,json=payload);assert response.status_code==200,response.text;return response.json()
def _get(client,path):
    response=client.get(path);assert response.status_code==200,response.text;return response.json()
@contextmanager
def _client():
    engine=create_engine("sqlite://",connect_args={"check_same_thread":False},poolclass=StaticPool);Base.metadata.create_all(engine);factory=sessionmaker(bind=engine,autoflush=False)
    def override():
        session=factory()
        try:yield session
        finally:session.close()
    app.dependency_overrides[get_db]=override
    try:
        with TestClient(app) as client:yield client
    finally:app.dependency_overrides.clear();Base.metadata.drop_all(engine);get_vector_store.cache_clear();get_embedding_service.cache_clear()
