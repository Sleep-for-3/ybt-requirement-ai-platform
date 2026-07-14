from io import BytesIO
from pathlib import Path
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from app.services.sql_parser import parse_sql
from .base import KnowledgeUnitDraft

QUESTION={"问题","答疑问题","监管问题","银行问题"};ANSWER={"回复","监管回复","答疑结果","处理意见"};SUGGESTION={"同业建议"};FIELD_CODE={"字段代码","数据项编码"};FIELD_NAME={"字段名称","字段中文名","数据项名称"};TABLE={"一表通表","表代码"};SCENARIO={"业务场景","产品场景","场景名称"};BUSINESS_SYSTEM={"来源系统","业务系统","系统名称"};SOURCE_TABLE={"来源表","来源表英文名","源表","源表英文名"};SOURCE_FIELD={"来源字段","来源字段英文名","源字段","源字段英文名"};MART_TABLE={"集市表","监管集市表"};MART_FIELD={"集市字段","监管集市字段"}
def parse_document(file_name,content,knowledge_type):
    suffix=Path(file_name).suffix.lower()
    if suffix not in {".xlsx",".docx",".pdf",".txt",".md",".sql"}:raise ValueError(f"Unsupported knowledge document format: {suffix or 'missing extension'}")
    if suffix==".xlsx":return _excel(content,knowledge_type)
    if suffix==".docx":return _docx(content)
    if suffix==".pdf":return _pdf(content)
    text=content.decode("utf-8",errors="replace")
    if suffix==".sql":
        parsed=parse_sql(text);summary=f"来源表: {', '.join(parsed.source_tables)}\n来源字段: {', '.join(parsed.selected_fields)}\n关联条件: {'; '.join(parsed.joins)}\n过滤条件: {'; '.join(parsed.where_conditions)}"
        return [KnowledgeUnitDraft("sql_summary",file_name,summary,metadata={"raw_sql":text,"parsed_success":parsed.parsed_success})],[parsed.error_message] if parsed.error_message else []
    return [KnowledgeUnitDraft("paragraph",file_name,part) for part in text.split("\n\n") if part.strip()],[]

def _excel(content,knowledge_type):
    wb=load_workbook(BytesIO(content),data_only=True);units=[];warnings=[]
    for sheet in wb.worksheets:
        matrix=_expanded_matrix(sheet)
        header_row=None;headers={}
        for row in range(1,min(len(matrix),50)+1):
            values={column:str(matrix[row-1][column-1] or "").strip() for column in range(1,sheet.max_column+1)}
            if any(value in QUESTION|ANSWER|FIELD_CODE|FIELD_NAME for value in values.values()):header_row=row;headers=values;break
        if not header_row:warnings.append(f"{sheet.title} 未识别表头");continue
        for row in range(header_row+1,sheet.max_row+1):
            values={headers[column]:matrix[row-1][column-1] for column in headers if headers[column]}
            if not any(value not in (None,"") for value in values.values()):continue
            question=next((str(values[key]) for key in QUESTION if values.get(key)),"");answer=next((str(values[key]) for key in ANSWER if values.get(key)),"");suggestion=next((str(values[key]) for key in SUGGESTION if values.get(key)),"")
            field_code=next((str(values[key]) for key in FIELD_CODE if values.get(key)),None);field_name=next((str(values[key]) for key in FIELD_NAME if values.get(key)),None);table=next((str(values[key]) for key in TABLE if values.get(key)),None);scenario_name=next((str(values[key]) for key in SCENARIO if values.get(key)),None);business_system=next((str(values[key]) for key in BUSINESS_SYSTEM if values.get(key)),None);source_table=next((str(values[key]) for key in SOURCE_TABLE if values.get(key)),None);source_field=next((str(values[key]) for key in SOURCE_FIELD if values.get(key)),None);mart_table=next((str(values[key]) for key in MART_TABLE if values.get(key)),None);mart_field=next((str(values[key]) for key in MART_FIELD if values.get(key)),None)
            content_text="\n".join(filter(None,[f"问题：{question}" if question else "",f"回复：{answer}" if answer else "",f"同业建议：{suggestion}" if suggestion else "", "；".join(f"{key}：{value}" for key,value in values.items() if value not in (None,"") and key not in QUESTION|ANSWER|SUGGESTION)]))
            unit_type="qa" if question or knowledge_type=="regulatory_qa" else "field_mapping" if field_code else "table_row"
            units.append(KnowledgeUnitDraft(unit_type,question or field_name or f"{sheet.title} 第{row}行",content_text,source_sheet_name=sheet.title,source_cell_range=f"A{row}:{get_column_letter(sheet.max_column)}{row}",target_table_code=table,target_field_code=field_code,target_field_name=field_name,scenario_name=scenario_name,source_table_name=source_table,source_field_name=source_field,metadata={"row":row,"scenario_name":scenario_name,"business_system_name":business_system,"mart_table_name":mart_table,"mart_field_name":mart_field}))
    return units,warnings

def _expanded_matrix(sheet):
    matrix=[[sheet.cell(row,column).value for column in range(1,sheet.max_column+1)] for row in range(1,sheet.max_row+1)]
    for merged in sheet.merged_cells.ranges:
        value=matrix[merged.min_row-1][merged.min_col-1]
        for row in range(merged.min_row,merged.max_row+1):
            for column in range(merged.min_col,merged.max_col+1):matrix[row-1][column-1]=value
    return matrix

def _docx(content):
    from docx import Document
    doc=Document(BytesIO(content));units=[];heading=None
    for index,p in enumerate(doc.paragraphs,1):
        text=p.text.strip()
        if not text:continue
        if p.style and p.style.name.lower().startswith("heading"):heading=text;continue
        units.append(KnowledgeUnitDraft("policy_clause",heading or f"段落 {index}",text,source_heading=heading,metadata={"paragraph":index}))
    for table_index,table in enumerate(doc.tables,1):
        for row_index,row in enumerate(table.rows,1):
            text=" | ".join(cell.text.strip() for cell in row.cells)
            if text.strip(" | "):units.append(KnowledgeUnitDraft("table_row",f"表格 {table_index} 第{row_index}行",text,source_heading=heading,source_cell_range=f"table:{table_index}:row:{row_index}"))
    return units,[]

def _pdf(content):
    from pypdf import PdfReader
    units=[];warnings=[]
    for page_no,page in enumerate(PdfReader(BytesIO(content)).pages,1):
        text=(page.extract_text() or "").strip()
        if text:units.append(KnowledgeUnitDraft("policy_clause",f"第 {page_no} 页",text,source_page_no=page_no))
        else:warnings.append(f"第 {page_no} 页无可提取文本，未启用 OCR")
    return units,warnings
